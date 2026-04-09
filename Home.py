"""
AI Contract Clause Builder - Assignment 2
（模块级说明；更细的分区见下方以 # 开头的「代码地图」）
"""

# =============================================================================
# 代码地图（后端复盘 / 学习用）
# -----------------------------------------------------------------------------
# 1) 配置与静态资源      — secrets_ai_contract_app/credentials.env、Frontview 密码页 HTML
# 2) 访问控制与 API Key   — 非 Agent；安全边界：密码与 Key 仅服务端
# 3) LLM 底座            — call_openai_chat / _openai_chat_message；标准对话与 tools 两种出口
# 4) 提示词与 Agent 编排  — ReAct 工具协议、各步 system/user、Reflexion Critic、Intake 标签化、记忆追问压缩上下文
# 5) RAG 与分块策略      — 文档级整文件为语料单元；检索用每文档前 800 字 + LLM 选编号或词重叠；注入时再截断
# 6) 文书引导（Intake）   — 无上传时合成语料池再 RAG；槽位追问与门禁
# 7) 文档解析            — extract_text：整文件为一条语料；供主工作流与 Intake 共用
# 8) Word 与主工作流      — Word 导出 + 顺序 DAG + 每步「ReAct → Critic → replan」
# 9) Streamlit 脚本入口   — 侧栏记忆、页面控件、门禁、compose objective 后调工作流
# =============================================================================

import streamlit as st
import streamlit.components.v1 as components
import os
import re
import sys
import json
import uuid
from pathlib import Path
from typing import Optional
from io import BytesIO
from datetime import datetime
from docx import Document
from openai import OpenAI
import traceback
from dotenv import load_dotenv

from rl_bandit import get_rerank_bandit
from vector_chunking import chunk_label, documents_to_vector_chunks

# --- 1) 配置与路径：与 Agent 无关，决定运行时可用的密钥与静态资源 ---
# 密钥唯一目录：secrets_ai_contract_app/credentials.env（复制 credentials.env.example）
_PROJECT_ROOT = Path(__file__).resolve().parent
_USER_MEMORY_ROOT = _PROJECT_ROOT / "user_memory"
_SECRETS_DIR = _PROJECT_ROOT / "secrets_ai_contract_app"
_APP_CREDENTIALS_ENV = _SECRETS_DIR / "credentials.env"
_SECRETS_FILE_HINT = "secrets_ai_contract_app/credentials.env"

_OPENAI_KEY_PLACEHOLDERS = frozenset(
    {
        "your_openai_api_key_here",
        "sk-your-openai-api-key-here",
    }
)

# region agent log (debug session ca4d61)
_DEBUG_LOG_PATH = _PROJECT_ROOT / ".cursor" / "debug-ca4d61.log"
# 镜像：*.log 常被工具链忽略；同内容 NDJSON 便于在本仓库内读取
_DEBUG_LOG_MIRROR_NDJSON = _PROJECT_ROOT / "debug-ca4d61.ndjson"
_DEBUG_INGEST_URL = "http://127.0.0.1:7414/ingest/84a485a2-dca9-4dca-8a8f-54f0d8fa96c5"


def _dbg_ca4d61(
    hypothesis_id: str,
    location: str,
    message: str,
    data: dict,
    *,
    run_id: str = "pre-fix",
) -> None:
    """NDJSON 调试日志；禁止写入密钥明文。"""
    line = json.dumps(
        {
            "sessionId": "ca4d61",
            "runId": run_id,
            "hypothesisId": hypothesis_id,
            "location": location,
            "message": message,
            "data": data,
            "timestamp": int(datetime.now().timestamp() * 1000),
        },
        ensure_ascii=False,
    )
    any_ok = False
    for p in (_DEBUG_LOG_PATH, _DEBUG_LOG_MIRROR_NDJSON):
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            with open(p, "a", encoding="utf-8") as _df:
                _df.write(line + "\n")
            any_ok = True
        except OSError:
            continue
    if not any_ok:
        try:
            import urllib.request

            req = urllib.request.Request(
                _DEBUG_INGEST_URL,
                data=line.encode("utf-8"),
                headers={
                    "Content-Type": "application/json",
                    "X-Debug-Session-Id": "ca4d61",
                },
                method="POST",
            )
            urllib.request.urlopen(req, timeout=0.35)
        except Exception:
            try:
                print(f"[debug-ca4d61] {line}", file=sys.stderr, flush=True)
            except OSError:
                pass


# endregion

if _APP_CREDENTIALS_ENV.exists():
    load_dotenv(_APP_CREDENTIALS_ENV, override=True)

# region agent log (debug session ca4d61)
_env_loaded_from = "credentials" if _APP_CREDENTIALS_ENV.exists() else "none"
_raw_openai = (os.getenv("OPENAI_API_KEY") or "").strip()
_dbg_ca4d61(
    "H1,H2",
    "Home.py:after_load_dotenv",
    "OPENAI_API_KEY from process env after load_dotenv",
    {
        "env_file_used": _env_loaded_from,
        "openai_key_len": len(_raw_openai),
        "openai_empty": len(_raw_openai) == 0,
        "openai_is_known_placeholder": _raw_openai in _OPENAI_KEY_PLACEHOLDERS,
        "openai_has_sk_prefix": _raw_openai.startswith("sk-") if _raw_openai else False,
    },
)
# endregion


def _normalize_gate_text(s: str) -> str:
    """统一换行、BOM、首尾空白，避免与 .env 肉眼一致却校验失败。"""
    if s is None:
        return ""
    t = str(s).replace("\r", "").replace("\ufeff", "")
    return t.strip()


def _strip_env_value_trailing_comment(val: str) -> str:
    """去掉行尾「空格 + # 注释」（密码里含 # 时请用引号包裹整段值）。"""
    val = val.strip()
    if " #" in val:
        val = val.split(" #", 1)[0].rstrip()
    return val


def _backend_password_from_file(env_path: Path) -> str:
    """直接从 .env 读取 BACKEND_PASSWORD；同一文件多次出现时以最后一行为准（覆盖误复制的多行）。"""
    if not env_path.is_file():
        return ""
    try:
        content = env_path.read_text(encoding="utf-8-sig")
    except OSError:
        return ""
    last: str = ""
    for raw in content.splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, _, val = line.partition("=")
        if key.strip() != "BACKEND_PASSWORD":
            continue
        val = _strip_env_value_trailing_comment(val)
        if len(val) >= 2 and val[0] == val[-1] and val[0] in "\"'":
            val = val[1:-1]
        last = _normalize_gate_text(val)
    return last


def _openai_api_key_from_file(env_path: Path) -> str:
    """直接从 .env 读取 OPENAI_API_KEY；多次出现时以最后一行为准（与 BACKEND_PASSWORD 解析一致）。"""
    if not env_path.is_file():
        return ""
    try:
        content = env_path.read_text(encoding="utf-8-sig")
    except OSError:
        return ""
    last: str = ""
    for raw in content.splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, _, val = line.partition("=")
        if key.strip() != "OPENAI_API_KEY":
            continue
        val = _strip_env_value_trailing_comment(val)
        if len(val) >= 2 and val[0] == val[-1] and val[0] in "\"'":
            val = val[1:-1]
        last = _normalize_gate_text(val)
    return last


def _env_openai_key_effective() -> str:
    """
    从 secrets_ai_contract_app/credentials.env 解析 OPENAI_API_KEY，
    避免 shell 或历史进程里残留的占位符覆盖「文件里已是 sk- 真钥」的情况。
    解析到有效 Key 时写回 os.environ，便于其它仅读 getenv 的代码路径一致。
    """
    for path in (_APP_CREDENTIALS_ENV,):
        k = _openai_api_key_from_file(path)
        if not k:
            continue
        if k in _OPENAI_KEY_PLACEHOLDERS:
            os.environ.pop("OPENAI_API_KEY", None)
            continue
        os.environ["OPENAI_API_KEY"] = k
        return k
    k = (os.getenv("OPENAI_API_KEY") or "").strip()
    if k in _OPENAI_KEY_PLACEHOLDERS:
        return ""
    return k


# 每轮 Streamlit 脚本重跑时从文件刷新进 os.environ，避免仅 shell 里残留占位符
_env_openai_key_effective()


def resolve_backend_password() -> str:
    """
    门禁密码：优先 credentials.env，其次环境变量 BACKEND_PASSWORD。
    """
    for path in (_APP_CREDENTIALS_ENV,):
        p = _backend_password_from_file(path)
        if p:
            return p
    return _normalize_gate_text(os.getenv("BACKEND_PASSWORD", ""))


# Frontview 密码页：components.html 嵌入 iframe，与 Streamlit 表单分工（眼睛动效 / 服务端校验密码）
_FRONTPAGE_EYES_HTML = _PROJECT_ROOT / "static" / "frontview_eyes.html"

# 仍为模板时提示用户改 secrets_ai_contract_app/credentials.env（避免误以为应用坏了）
_KNOWN_PLACEHOLDER_PASSWORDS = frozenset(
    {
        "your_backend_password_here",
        "your-login-password-here",
        "your_login_password_here",
    }
)


# --- 2) 访问控制与 API Key 解析（非 LLM Agent）---
# 知识：部署模式分三种 — (a) BACKEND_PASSWORD + OPENAI_API_KEY 托管 (b) 仅环境变量 Key (c) 用户自填 Key
def get_api_key_main():
    """
    Resolve OpenAI API key in main area (no sidebar).
    After correct BACKEND_PASSWORD, uses server OPENAI_API_KEY from credentials.env automatically.
    """
    _raw_openai_env = (os.getenv("OPENAI_API_KEY") or "").strip()
    _raw_openai_file = _openai_api_key_from_file(_APP_CREDENTIALS_ENV)
    server_key = _env_openai_key_effective()
    pwd_required = resolve_backend_password()
    verified = st.session_state.get("backend_verified")

    # region agent log (debug session ca4d61)
    _dbg_ca4d61(
        "H1,H3,H5",
        "get_api_key_main:entry",
        "resolver inputs",
        {
            "verified": bool(verified),
            "pwd_required_non_empty": bool(pwd_required),
            "raw_env_openai_len": len(_raw_openai_env),
            "raw_env_is_placeholder": _raw_openai_env in _OPENAI_KEY_PLACEHOLDERS,
            "file_parsed_openai_len": len(_raw_openai_file),
            "file_parsed_is_placeholder": _raw_openai_file in _OPENAI_KEY_PLACEHOLDERS,
            "effective_server_key_len": len(server_key),
            "widget_api_key_len": len(
                (st.session_state.get("api_key_input") or "").strip()
            ),
        },
    )
    # endregion

    # Password verified: prefer server key so user needs no own key
    if verified and server_key:
        # region agent log (debug session ca4d61)
        _dbg_ca4d61(
            "H1,H3",
            "get_api_key_main:branch",
            "return verified_and_server_key",
            {"branch": "verified_and_server_key", "return_len": len(server_key)},
        )
        # endregion
        st.success(
            "已通过访问验证 · 系统已自动使用服务端 API Key，您可直接跳转使用合同拟制助手。"
        )
        return server_key

    if verified and pwd_required and not server_key:
        # region agent log (debug session ca4d61)
        _dbg_ca4d61(
            "H1",
            "get_api_key_main:branch",
            "verified_but_no_server_key",
            {
                "branch": "verified_no_effective_openai",
                "raw_was_placeholder": _raw_openai_env in _OPENAI_KEY_PLACEHOLDERS
                or _raw_openai_file in _OPENAI_KEY_PLACEHOLDERS,
            },
            run_id="post-fix",
        )
        # endregion
        if (
            _raw_openai_env in _OPENAI_KEY_PLACEHOLDERS
            or _raw_openai_file in _OPENAI_KEY_PLACEHOLDERS
        ):
            st.error(
                "已通过验证，但 **OPENAI_API_KEY** 在配置文件中仍是**示例占位符**（如 your_openai_api_key_here）。"
                f"请在 `{_SECRETS_FILE_HINT}` 中写入以 **sk-** 开头的真实密钥，**保存文件**后**重启 Streamlit**；"
                "密钥中间**不要有空格**。您也可在下方「API 配置」中临时粘贴 Key。"
            )
        else:
            st.error(
                "已通过验证，但未读取到有效的 OPENAI_API_KEY。请在「API 配置」中填写您自己的 Key，"
                f"或由管理员在 {_SECRETS_FILE_HINT} 中配置 OPENAI_API_KEY。"
            )

    # Deployed with key in env but no access password: use env key for everyone
    if server_key and not pwd_required:
        # region agent log (debug session ca4d61)
        _dbg_ca4d61(
            "H2,H3",
            "get_api_key_main:branch",
            "no_password_gate_use_env_key",
            {"branch": "server_key_no_pwd_required"},
        )
        # endregion
        st.info("当前部署已启用环境变量中的 API Key，可直接使用。")
        return server_key

    expand = not server_key
    with st.expander("API 配置（未自动使用服务端 Key 时在此填写）", expanded=expand):
        key = st.text_input(
            "OpenAI API Key",
            type="password",
            key="api_key_input",
            help="https://platform.openai.com/api-keys",
        )
        if key:
            st.success("API Key 已填写，可启动助手。")
        else:
            st.warning("请填写 API Key，或完成上方访问验证以使用服务端 Key。")
        # region agent log (debug session ca4d61)
        _out = (key or "").strip() or None
        if _out and _out in _OPENAI_KEY_PLACEHOLDERS:
            _out = None
        _dbg_ca4d61(
            "H3,H5",
            "get_api_key_main:branch",
            "return manual expander key",
            {
                "branch": "manual_api_key_input",
                "return_is_none": _out is None,
                "return_len": len(_out) if _out else 0,
            },
        )
        # endregion
        return _out


# render_access_login_page：前端门禁 UI；Agent 知识不适用（纯鉴权与 UX）
def render_access_login_page(app_title: str, backend_password: str) -> None:
    """
    Password gate UI adapted from https://github.com/Umia888/Frontviewpage (Eye + EyeCard layout).
    Eyes run in an iframe component; password is verified only on the Streamlit server.
    """
    eye_state = "wrong" if st.session_state.pop("_pwd_gate_show_wrong_eyes", False) else "normal"
    if st.session_state.pop("_pwd_gate_flash_error", False):
        st.error("密码错误，请重试。")

    st.markdown(
        """
<style>
[data-testid="stAppViewContainer"] {
  background-color: #0062ad !important;
}
[data-testid="stHeader"] { background-color: transparent !important; }
section.main > div { padding-top: 0.5rem !important; }
label[data-testid="stWidgetLabel"] p { color: #ffffff !important; }
</style>
""",
        unsafe_allow_html=True,
    )

    col_l, col_r = st.columns([1.08, 1.0], gap="large")
    with col_l:
        if _FRONTPAGE_EYES_HTML.is_file():
            tpl = _FRONTPAGE_EYES_HTML.read_text(encoding="utf-8")
            eyes_html = tpl.replace("__EYE_STATE__", json.dumps(eye_state))
            components.html(eyes_html, height=460, scrolling=False)
        else:
            st.warning("未找到 static/frontview_eyes.html，使用备用布局。")
            st.caption("请从仓库恢复 Frontview 眼睛组件静态页。")

    with col_r:
        st.markdown(
            f"""
<div style="color:#ffffff;">
  <p style="font-size:0.8rem;opacity:0.85;margin:0 0 0.35rem 0;">{app_title}</p>
  <p style="font-size:1.5rem;font-weight:700;margin:0 0 0.5rem 0;line-height:1.25;">请在下方输入登录密码</p>
  <p style="font-size:0.95rem;opacity:0.95;margin:0 0 1rem 0;">验证成功后，系统将<strong>自动使用服务器配置的 OpenAI API Key</strong>，无需自行填写 Key。</p>
</div>
""",
            unsafe_allow_html=True,
        )
        _exp_preview = _normalize_gate_text(backend_password)
        if _exp_preview in _KNOWN_PLACEHOLDER_PASSWORDS:
            st.error(
                "检测到 **BACKEND_PASSWORD** 仍是示例占位符。请编辑 `secrets_ai_contract_app/credentials.env`，"
                "把该项改成您自己的密码并保存，然后刷新本页。"
            )
        # 使用 st.form：避免「点按钮的同一轮 rerun」里 text_input 返回值未带上密码，导致永远校验失败。
        with st.form("access_gate_form", clear_on_submit=False):
            pwd = st.text_input(
                "访问密码",
                type="password",
                key="access_gate_password",
                autocomplete="current-password",
                placeholder="输入密码",
                label_visibility="collapsed",
            )
            go = st.form_submit_button("验证并进入", type="primary", use_container_width=True)
        st.markdown(
            """
<p style="color:#000000;font-size:14px;margin-top:1.25rem;max-width:28rem;line-height:1.45;
background:#fbf0dc;padding:12px 14px;border-radius:12px;">
忘记密码？请联系管理员或在 <code>secrets_ai_contract_app/credentials.env</code> 中查看 <code>BACKEND_PASSWORD</code> 配置说明。
</p>
""",
            unsafe_allow_html=True,
        )
        if go:
            _inp = _normalize_gate_text(
                st.session_state.get("access_gate_password") or pwd or ""
            )
            _exp = _normalize_gate_text(resolve_backend_password())
            if _exp and _inp == _exp:
                st.session_state.backend_verified = True
                st.success("验证成功，正在进入…")
                st.rerun()
            else:
                st.session_state._pwd_gate_show_wrong_eyes = True
                st.session_state._pwd_gate_flash_error = True
                st.rerun()

    st.markdown(
        """
<div style="text-align:center;margin-top:1.5rem;color:#ffffff;font-size:10px;opacity:0.88;">
  交互动效与视觉参考
  <a href="https://github.com/Umia888/Frontviewpage" target="_blank" rel="noopener noreferrer"
     style="color:#ffffff;text-decoration:underline;">Frontviewpage</a>
  · 致谢 Rafael Serra（原设计）
</div>
<div style="text-align:center;margin-top:0.35rem;">
  <a href="https://anotherplanet.io/" target="_blank" rel="noopener noreferrer"
     style="color:#ffffff;font-size:10px;opacity:0.88;">anotherplanet.io</a>
</div>
""",
        unsafe_allow_html=True,
    )


# --- 3) LLM 底座：标准 Chat Completions（无 tools）---
# 【提示词工程】messages 由调用方组装：system 定角色/约束，user 放任务与材料；本函数不负责写死业务提示词。
# 【Agent 调用工程】所有「单次生成」子调用（如 brainstorm_bullets、检索排序、Intake）统一走此出口，便于改模型与计次。
# 【参数】temperature 低 → 合同/检索排序更稳；max_tokens 在调用点按步骤收紧，控制成本与截断风险。
# 【观测】st.session_state.ai_call_count 按请求次数累加（非 Token 精确计量）。
def call_openai_chat(messages, api_key, model="gpt-4o-mini", temperature=0.2, max_tokens=2000):
    """OpenAI Chat Completions；失败时 st.error + st.stop 阻断流水线。"""
    try:
        if 'ai_call_count' not in st.session_state:
            st.session_state.ai_call_count = 0
        st.session_state.ai_call_count += 1
        
        # OpenAI API Configuration
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.openai.com/v1"
        )
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            timeout=60  # 60
        )
        
        return resp.choices[0].message.content.strip()
        
    except Exception as e:
        error_type = type(e).__name__
        # region agent log (debug session ca4d61)
        _dbg_ca4d61(
            "H4,H5",
            "call_openai_chat:except",
            "openai request failed",
            {
                "error_type": error_type,
                "api_key_len": len((api_key or "")),
                "api_key_missing": not (api_key or "").strip(),
            },
        )
        # endregion
        
        # 
        if "AuthenticationError" in error_type:
            st.error("**API Authentication Failed**\n\nPlease check your OpenAI API key.")
        elif "RateLimitError" in error_type:
            st.error(" **API**\n\n\n- \n- API\n- ")
        elif "timeout" in str(e).lower():
            st.error(" ****\n\nAPI")
        else:
            st.error(f" **API**\n\n: {error_type}\n: {str(e)}")
        
        # 
        with st.expander(" View Detailed Error InformationFor debugging"):
            st.code(traceback.format_exc())
        
        st.stop()


# --- 3b) ReAct 工具循环：Chat Completions + tools（Thought → Action → Observation）---
# 【Agent 调用工程】把「一步流水线生成」拆成多轮：模型选 function name + JSON 参数 → 本地执行 → tool 消息回灌。
# 【提示词工程】工具 schema 的 name/description/parameters 即对模型的「可行动作说明」，影响何时调用何工具。
# 【与 Reflexion 关系】ReAct 管步内推理与检索；外层 agent_call_with_verify_and_replan 再用 Critic + meta-prompt 做步外重试。
# 四工具分工：think_aloud=显式推理占位；brainstorm_bullets=嵌套子调用扩写；rag_retrieve=对接语料 RAG；submit_final_answer=结束工具环并交卷。
_REACT_TOOL_SPECS = [
    {
        "type": "function",
        "function": {
            # 工具 1：思维链外置，减少最终答案里的冗长推导（Observation 仍占上下文，故 think_aloud 结果在 handler 内也截断）。
            "name": "think_aloud",
            "description": "Record intermediate reasoning (chain-of-thought scratchpad). Does not change external state.",
            "parameters": {
                "type": "object",
                "properties": {
                    "note": {"type": "string", "description": "Your current reasoning step or hypothesis"},
                },
                "required": ["note"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            # 工具 2：对窄问题再调一次 chat，实现「主 Agent 调度 + 子生成器」模式。
            "name": "brainstorm_bullets",
            "description": "Ask the LLM for short bullet ideas on a focused sub-question (sub-agent call).",
            "parameters": {
                "type": "object",
                "properties": {
                    "focus_question": {
                        "type": "string",
                        "description": "A narrow question to explore before drafting the final answer",
                    },
                },
                "required": ["focus_question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            # 工具 3：参数 top_k 在 schema 里描述为 Max chunks，实为「最多返回文档条数」而非向量切片数。
            "name": "rag_retrieve",
            "description": "Retrieve relevant excerpts from uploaded documents (RAG). If no uploads, returns a notice.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string"},
                    "top_k": {"type": "integer", "description": "Max chunks", "default": 3},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            # 工具 4：与 react_addon 文案一致，形成「必须显式交卷」的停止条件，避免无限工具循环。
            "name": "submit_final_answer",
            "description": "Submit the complete deliverable for this pipeline step. Call this once when ready.",
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {"type": "string", "description": "Full final output for this step"},
                },
                "required": ["content"],
            },
        },
    },
]


def _openai_chat_message(
    api_key,
    messages,
    *,
    model="gpt-4o-mini",
    temperature=0.2,
    max_tokens=2000,
    tools=None,
    tool_choice="auto",
):
    # 【Agent 调用工程】与 call_openai_chat 对齐，但返回完整 assistant message（含 tool_calls），供 ReAct 循环拼接多轮。
    # 【参数】tool_choice="auto"：由模型决定是否调用工具；若改为 required 可强制工具化（本项未用）。
    try:
        if "ai_call_count" not in st.session_state:
            st.session_state.ai_call_count = 0
        st.session_state.ai_call_count += 1
        client = OpenAI(api_key=api_key, base_url="https://api.openai.com/v1")
        kwargs = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "timeout": 60,
        }
        if tools:
            kwargs["tools"] = tools
            kwargs["tool_choice"] = tool_choice
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message
    except Exception as e:
        error_type = type(e).__name__
        if "AuthenticationError" in error_type:
            st.error("**API Authentication Failed**\n\nPlease check your OpenAI API key.")
        elif "RateLimitError" in error_type:
            st.error("API Rate limit / quota issue.")
        elif "timeout" in str(e).lower():
            st.error("API timeout.")
        else:
            st.error(f"API error: {error_type}: {str(e)}")
        with st.expander("View Detailed Error Information"):
            st.code(traceback.format_exc())
        st.stop()


def _assistant_message_to_dict(msg) -> dict:
    d = {"role": "assistant", "content": msg.content if msg.content else None}
    tcs = getattr(msg, "tool_calls", None)
    if tcs:
        d["tool_calls"] = [
            {
                "id": tc.id,
                "type": "function",
                "function": {
                    "name": tc.function.name,
                    "arguments": tc.function.arguments or "{}",
                },
            }
            for tc in tcs
        ]
    return d


def _build_react_tool_handlers(ctx: dict, api_key: str) -> dict:
    # 【Agent 调用工程】工具名 → Python 可调用对象；返回值字符串即 ReAct 的 Observation，回写给模型继续推理。
    # 【RAG】rag_retrieve 内复用 ai_enhanced_retrieve / simple_retrieve，与主流程同一套检索逻辑。

    def think_aloud(note: str, **_) -> str:
        # 【提示词工程】不显式调用 LLM；把「思考」外化为可记录事件，减轻主答案里夹带长链式思考。
        # 【分块】note 截断 1200 字，防止单次 Observation 撑爆上下文。
        return f"[think_aloud] Recorded: {note[:1200]}"

    def brainstorm_bullets(focus_question: str, **_) -> str:
        # 【Agent 调用工程】子 Agent：一次独立 chat，专门扩写要点，与主步 final answer 解耦。
        # 【提示词工程】system 强制「短 bullet」，user 仅为子问题，属任务分解（decomposition）。
        return call_openai_chat(
            [
                {
                    "role": "system",
                    "content": "You are a concise legal drafting assistant. Reply in short bullets only.",
                },
                {"role": "user", "content": focus_question},
            ],
            api_key,
            temperature=0.25,
            max_tokens=700,
        )

    def rag_retrieve(query: str, top_k: int = 3, **_) -> str:
        # 【RAG 调用工程】query 由模型根据当前推理生成；top_k 上限钳制到 5，平衡证据量与 token。
        # 【向量切片】Observation 注入命中块的前 1000 字；语料已为 chunk 列表（见 documents_to_vector_chunks）。
        texts = ctx.get("texts") or []
        if not texts:
            return "Observation: no uploaded documents in corpus; proceed without RAG or use brainstorm_bullets."
        k = max(1, min(int(top_k), 5))
        r = ai_enhanced_retrieve(texts, query, api_key, top_k=k)
        if not r:
            r = simple_retrieve(texts, query, top_k=k)
        if not r:
            return "Observation: RAG returned no chunks for this query."
        parts = [f"== {chunk_label(x)} ==\n{x['text'][:1000]}" for x in r]
        return "\n\n".join(parts)

    def submit_final_answer(content: str, **_) -> str:
        # 【Agent 调用工程】显式终止工具循环：把最终交付物写入 ctx，供 run_react_tool_loop 提前返回。
        ctx["_react_final"] = content
        return "Observation: final answer stored. Do not call tools again; end turn."

    return {
        "think_aloud": think_aloud,
        "brainstorm_bullets": brainstorm_bullets,
        "rag_retrieve": rag_retrieve,
        "submit_final_answer": submit_final_answer,
    }


def run_react_tool_loop(
    api_key: str,
    *,
    system_prompt: str,
    user_message: str,
    react_ctx: dict,
    max_tool_turns: int = 16,
    model: str = "gpt-4o-mini",
    temperature: float = 0.2,
    max_tokens: int = 2000,
) -> tuple:
    """
    完整 ReAct：反复 assistant(tool_calls) → tool(observation) 直至无 tool_calls 或 submit_final_answer。
    返回 (final_text, trace)；trace 为 list[{tool, args, observation}] 按时间顺序。
    """
    # 【提示词工程】在业务 system_prompt 后拼接「工具协议」：规定终止条件（submit_final_answer）与可用动作，属典型 ReAct 指令模板。
    react_addon = (
        "\n\n[ReAct — tool protocol]\n"
        "You must eventually call submit_final_answer(content=...) with your **complete** deliverable for the user task.\n"
        "Until then you may call: think_aloud(note), brainstorm_bullets(focus_question), rag_retrieve(query[, top_k]).\n"
        "After each tool result (Observation), continue reasoning and choose the next action until you submit the final answer."
    )
    messages: list = [
        {"role": "system", "content": system_prompt + react_addon},
        {"role": "user", "content": user_message},
    ]
    handlers = _build_react_tool_handlers(react_ctx, api_key)
    trace = []
    last_natural = ""
    react_ctx["_react_final"] = None

    for _ in range(max_tool_turns):
        msg = _openai_chat_message(
            api_key,
            messages,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            tools=_REACT_TOOL_SPECS,
            tool_choice="auto",
        )
        tcs = getattr(msg, "tool_calls", None)
        if tcs:
            messages.append(_assistant_message_to_dict(msg))
            for tc in tcs:
                name = tc.function.name
                raw_args = tc.function.arguments or "{}"
                try:
                    args = json.loads(raw_args)
                except json.JSONDecodeError:
                    args = {}
                fn = handlers.get(name)
                if not fn:
                    obs = f"Unknown tool: {name}"
                else:
                    try:
                        obs = fn(**args)
                    except TypeError as ex:
                        obs = f"Tool argument error for {name}: {ex}"
                # 【分块 / 预算】trace 与回传模型的 tool content 截断，避免单轮 Observation 过长拖垮后续生成。
                trace.append({"tool": name, "args": args, "observation": str(obs)[:2500]})
                messages.append(
                    {"role": "tool", "tool_call_id": tc.id, "content": str(obs)[:8000]}
                )
            if react_ctx.get("_react_final"):
                return str(react_ctx["_react_final"]).strip(), trace
            continue

        last_natural = (msg.content or "").strip()
        if last_natural:
            return last_natural, trace
        messages.append(_assistant_message_to_dict(msg))
        # 【提示词工程】模型若返回空内容且无 tool_calls，用 user 消息轻推，避免死循环（非替换 system，以免覆盖业务角色）。
        messages.append(
            {
                "role": "user",
                "content": "Provide either tool calls or a substantive final answer, or call submit_final_answer(content=...).",
            }
        )

    final = react_ctx.get("_react_final")
    if final and str(final).strip():
        return str(final).strip(), trace
    return last_natural, trace


# --- 4) 反思-重规划（Reflexion 风格）与固定交付物格式 ---
# 【Agent 调用工程】步内 ReAct 生成后，由独立「评判 Agent」做 PASS/FAIL；FAIL 则将 keywords + replan 注入下一轮 user（meta-prompt）。
# 【提示词工程】Critic 使用强结构标签（[VERDICT] 等）便于正则解析，避免自然语言含糊导致无法自动分支。
def regex_clean_agent_text(text: str) -> str:
    # 【提示词后处理】非向模型再提问；用规则削弱 markdown/噪声，便于固定版式交付与下游解析（与生成提示词互补）。
    if not text:
        return ""
    s = text.strip()
    s = re.sub(r"\r\n?", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", s)
    s = re.sub(r"^#{1,6}\s+", "", s, flags=re.M)
    s = re.sub(r"```[a-zA-Z0-9]*\n?", "", s)
    s = s.replace("```", "")
    s = re.sub(r"^\s{0,3}[*\-]\s+", "• ", s, flags=re.M)
    return s.strip()


def format_fixed_deliverable(step_key: str, cleaned_body: str, status: str = "VERIFIED") -> str:
    # 【提示词后处理 / 交付协议】固定头尾与 STEP/TIMESTAMP/STATUS，便于日志、导出或与外部系统对接（非自然语言契约）。
    """Fixed layout document after regex cleanup (machine-readable block)."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return (
        "=== CONTRACT_AGENT_DELIVERABLE ===\n"
        f"STEP: {step_key}\n"
        f"TIMESTAMP: {ts}\n"
        f"STATUS: {status}\n"
        "=== CONTENT ===\n"
        f"{cleaned_body}\n"
        "=== END ==="
    )


def _parse_reflection_verdict(reflection_text: str) -> dict:
    # 【提示词工程配套】将 Critic 的结构化输出解析为 verdict / keywords / replan，供代码驱动 replan 循环。
    t = reflection_text or ""
    vm = re.search(r"\[VERDICT\]\s*(PASS|FAIL)\b", t, re.I)
    verdict = vm.group(1).upper() if vm else "FAIL"

    kw_m = re.search(
        r"\[KEYWORDS_FOR_REPLAN\]\s*(.*?)(?=\[REPLAN_HINT\]|\Z)", t, re.S | re.I
    )
    keywords = (kw_m.group(1).strip() if kw_m else "").strip() or "clarity, completeness, feasibility"

    rh_m = re.search(r"\[REPLAN_HINT\]\s*(.*)\Z", t, re.S | re.I)
    replan = (rh_m.group(1).strip() if rh_m else "").strip()
    if not replan:
        replan = (
            "Tighten alignment with the task, remove hedging where inappropriate, "
            "and ensure all requested sections are explicitly answered."
        )
    return {"verdict": verdict, "keywords": keywords, "replan": replan, "raw": t}


def _run_reflection_critic(
    api_key,
    *,
    agent_label: str,
    task_summary: str,
    model_output: str,
    context_snapshot: str,
) -> dict:
    # 【Agent 调用工程】单步单次评判，不调用工具；temperature 极低以减少评分抖动。
    # 【提示词工程】明确要求三维标准（correct/complete/feasible）+ 固定标签块，便于自动化与下游 meta-prompt。
    # 【分块】model_output 预览上限 14000 字符，控制 Critic 上下文长度。
    out_preview = (model_output or "")[:14000]
    user = f"""You are a strict output critic for an LLM agent pipeline. Judge the draft below.

Agent step: {agent_label}
Task (what was asked): {task_summary}
Context (brief): {context_snapshot}

Model output:
---
{out_preview}
---

Decide if the output is simultaneously:
- **correct** (no major factual/logic errors relative to the task),
- **complete** (addresses the task requirements, no large missing parts),
- **feasible** (actionable / usable for downstream contract drafting steps).

Reply in EXACTLY this structure (plain text, no code fences):

[VERDICT]
PASS
or
FAIL

[CHECKLIST]
correct: YES or NO
complete: YES or NO
feasible: YES or NO

[ISSUES]
- (bullet points; or single line "none")

[KEYWORDS_FOR_REPLAN]
comma-separated 3–8 keywords or short phrases to refocus the next LLM call

[REPLAN_HINT]
One short paragraph: meta-instruction for rewriting the prompt / emphasis for the next attempt (what to fix).
"""
    raw = call_openai_chat(
        [
            # 【提示词工程】极简 system：只强调「精确 + 只输出标签块」，减少与 user 里长指令打架。
            {"role": "system", "content": "You are a precise critic. Output only the requested labeled sections."},
            {"role": "user", "content": user},
        ],
        api_key,
        temperature=0.1,
        max_tokens=700,
    )
    return _parse_reflection_verdict(raw)


def agent_call_with_verify_and_replan(
    api_key,
    *,
    step_key: str,
    step_label: str,
    task_summary: str,
    reflect_context: str,
    system_prompt: str,
    build_user_content,
    react_ctx: dict,
    max_tokens=1200,
    temperature=0.2,
    max_replans=2,
    max_react_turns: int = 16,
):
    """ReAct 生成 + Critic + meta replan；返回清洗正文、固定格式块、原始输出、反思记录、replan 次数。"""
    # 【Agent 调用工程】单流水线步的完整包裹：ReAct(max_react_turns) → Critic → 最多 max_replans 次外层 replan。
    # 【提示词工程】build_user_content(meta_block) 将业务 user 与 meta-prompt 拼接；首轮 meta_block 为空。
    # 【RAG】react_ctx["texts"] 传入 ReAct 的 rag_retrieve；无上传时传空列表则工具返回「无文档」Observation。
    meta_block = ""
    reflection_records = []
    last_raw = ""
    replan_count = 0

    for attempt in range(max_replans + 1):
        user_content = build_user_content(meta_block)
        ctx = dict(react_ctx)
        last_raw, react_trace = run_react_tool_loop(
            api_key,
            system_prompt=system_prompt,
            user_message=user_content,
            react_ctx=ctx,
            max_tool_turns=max_react_turns,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        if not (last_raw or "").strip():
            last_raw = (ctx.get("_react_final") or "").strip()

        critic = _run_reflection_critic(
            api_key,
            agent_label=f"{step_label} (attempt {attempt + 1})",
            task_summary=task_summary,
            model_output=last_raw,
            context_snapshot=reflect_context,
        )
        reflection_records.append(
            {
                "attempt": attempt + 1,
                "verdict": critic["verdict"],
                "raw_reflection": critic["raw"],
                "react_trace": react_trace,
            }
        )

        # 【RL / 多臂老虎机】收集本步 Critic 二元结果，供主工作流结束后对「检索排序臂」做奖励更新。
        if st.session_state.get("_rl_collect_verdicts"):
            st.session_state.setdefault("_rl_episode_verdicts", []).append(
                1.0 if critic["verdict"] == "PASS" else 0.0
            )

        if critic["verdict"] == "PASS":
            cleaned = regex_clean_agent_text(last_raw)
            fixed = format_fixed_deliverable(step_key, cleaned, status="VERIFIED")
            return cleaned, fixed, last_raw, reflection_records, replan_count

        if attempt >= max_replans:
            cleaned = regex_clean_agent_text(last_raw)
            fixed = format_fixed_deliverable(step_key, cleaned, status="UNVERIFIED_MAX_RETRIES")
            return cleaned, fixed, last_raw, reflection_records, replan_count

        replan_count += 1
        # 【提示词工程】meta-prompt：告诉模型上一轮未通过自动评审，并给出关键词 + 改写指令 + 负例摘录（防复读错误）。
        meta_block = (
            f"[META-PROMPT — REPLAN ATTEMPT {attempt + 2}]\n"
            f"The previous answer failed automated review (correct/complete/feasible).\n"
            f"Refocus keywords: {critic['keywords']}\n\n"
            f"Meta instruction for this attempt:\n{critic['replan']}\n\n"
            f"Previous output (for reference only, do not copy errors):\n"
            f"---\n{last_raw[:3500]}\n---\n"
            f"Use the ReAct tools again if needed, then submit_final_answer with an improved full deliverable."
        )

    cleaned = regex_clean_agent_text(last_raw)
    fixed = format_fixed_deliverable(step_key, cleaned, status="UNVERIFIED")
    return cleaned, fixed, last_raw, reflection_records, replan_count


def streamlit_show_agent_verification(
    step_label: str,
    cleaned: str,
    fixed: str,
    reflection_records: list,
    replan_count: int,
    main_render,
):
    """Show critic log + main body + regex-cleaned fixed-format block."""
    with st.expander(f"{step_label} · 验证与反思", expanded=replan_count > 0):
        for rec in reflection_records:
            st.markdown(f"**轮次 {rec['attempt']} · 结论：`{rec['verdict']}`**")
            st.markdown(rec["raw_reflection"])
            rt = rec.get("react_trace")
            if rt:
                with st.expander(
                    f"ReAct 工具轨迹（第 {rec['attempt']} 次尝试）", expanded=False
                ):
                    for j, ev in enumerate(rt, 1):
                        st.markdown(f"**{j}. `{ev.get('tool')}`**")
                        st.caption(
                            json.dumps(ev.get("args"), ensure_ascii=False)[:1200]
                            or "{}"
                        )
                        st.text((ev.get("observation") or "")[:4000])
        if replan_count > 0:
            st.info(f"本步骤已根据 meta-prompt 与关键词重选重新生成，共 replan {replan_count} 次。")
    main_render(cleaned)
    with st.expander(f"{step_label} · 固定格式文档（正则清洗）", expanded=False):
        st.code(fixed, language="text")


# =============================================================================
# 5) 文书拟制计划 + RAG 槽位引导（Intake / 前置「引导类 Agent」）
# -----------------------------------------------------------------------------
# 【Agent 工程】多角色合一：Planner（PLAN）+ 槽位提示（RAG_SLOT_HINTS）+ 追问（QUESTIONS）+ 状态机（HEURISTIC_STATUS）。
# 【提示词工程】英文方括号标签 + 「标签独占一行」约束，配合 _intake_tag_block 正则抽取，属结构化输出（cheap schema）。
# 【RAG】build_intake_rag_context_from_texts：先向量切片语料（见 vector_chunking.py），再检索；无上传则合成语料池后同样分块。
# 【控制流】忽略建议 / 跳过引导 / ready 门禁见 intake_workflow_gate_ok。
# =============================================================================
INTAKE_MAX_GUIDE_ROUNDS = 6


def _intake_tag_block(text: str, tag: str) -> str:
    # 【提示词工程配套】从模型输出中切分标签块；依赖模型遵守「[TAG] 下一行起为内容」的约定。
    if not text:
        return ""
    pat = rf"\[{re.escape(tag)}\]\s*(.*?)(?=\n\s*\[[A-Z0-9_]+\]|\Z)"
    m = re.search(pat, text, flags=re.DOTALL | re.IGNORECASE)
    return m.group(1).strip() if m else ""


def _intake_status_is_complete(status_blob: str) -> bool:
    t = (status_blob or "").strip().upper()
    if not t:
        return False
    if "COMPLETE" in t and "CONTINUE" not in t:
        return True
    if "CONTINUE" in t and "COMPLETE" not in t:
        return False
    if "COMPLETE" in t and "CONTINUE" in t:
        return t.rfind("COMPLETE") > t.rfind("CONTINUE")
    return False


def build_intake_rag_context_from_texts(texts, objective, jurisdiction, api_key) -> str:
    # 【RAG 调用工程】查询 q = objective + jurisdiction；top_k=4 控制注入条数。
    # 【向量切片】上传文件：.md/.txt 用递归字符分块；.docx 用语义分块（段落嵌入合并）；合成语料池为省成本强制递归分块。
    # 【检索后截断】每条命中 chunk 取 text[:1200] 拼进 planner 的 user，控制上下文长度。
    q = f"{objective}\n{jurisdiction}"
    had_uploads = bool(texts)
    corpus = list(texts) if texts else []

    if not corpus:
        # 【提示词工程】合成语料：要求分节标题 + 字数区间，便于后续「按段」被 simple/LLM 检索命中（仍是一条大 text）。
        synth = call_openai_chat(
            [
                {
                    "role": "system",
                    "content": "你是法律检索语料撰写助手，输出可供段落检索的参考正文，不构成法律意见。",
                },
                {
                    "role": "user",
                    "content": f"""法域：{jurisdiction or '未指定'}
用户诉求：{objective[:2500]}

请用中文撰写一份**类案/场景参考摘要稿**（约 1200–2000 字），分 4–6 个小节，每节有标题（如「一、…」「二、…」）。须包含：
- 类似商业或法律场景描述
- 常见合同条款关注点与待澄清槽位
- 该法域下常见的程序或实体提示（标注为常识性、非法条引用）

便于后续按主题检索；不要使用 Markdown 代码块。""",
                },
            ],
            api_key,
            temperature=0.25,
            max_tokens=2800,
        )
        corpus = [{"filename": "类案联想语料池.txt", "text": synth}]

    force_recursive_chunks = not had_uploads
    corpus_chunks = documents_to_vector_chunks(
        corpus, api_key, force_recursive=force_recursive_chunks
    )

    retrieved = ai_enhanced_retrieve(corpus_chunks, q, api_key, top_k=4)
    if not retrieved:
        retrieved = simple_retrieve(corpus_chunks, q, top_k=4)
    # 【RAG 兜底】仍无命中时退回第一个 chunk，保证 planner 至少有一点材料。
    if not retrieved and corpus_chunks:
        retrieved = corpus_chunks[:1]

    parts = []
    for r in retrieved:
        parts.append(f"【{chunk_label(r)}】\n{r['text'][:1200]}")

    tag = "用户上传文件经 RAG 检索" if had_uploads else "模型生成语料池经 RAG 检索"
    return f"（{tag}）\n\n" + "\n\n---\n\n".join(parts)


def generate_intake_plan_and_guidance(
    api_key,
    objective: str,
    jurisdiction: str,
    uploaded_files,
    preloaded_texts: Optional[list] = None,
) -> dict:
    """首轮：文书拟制 plan + RAG 槽位 + 问题 + 忽略说明。"""
    # 【RAG 语料来源】优先当前上传解析；否则用记忆中恢复的 preloaded_texts（与主工作流一致）。
    if uploaded_files:
        texts = extract_text_from_uploaded_files(uploaded_files)
    elif preloaded_texts:
        texts = [dict(x) for x in preloaded_texts]
    else:
        texts = []
    rag_ctx = build_intake_rag_context_from_texts(texts, objective, jurisdiction, api_key)
    # 【提示词工程】将 RAG 结果作为「证据段落」嵌入 user；标签块驱动 planner+槽位+追问+继续/结束信号。
    user = f"""你是资深合同与法律文书辅助顾问。用户的法律诉求与法域如下（表述可能不清晰）。

【用户诉求】
{objective}

【法域】
{jurisdiction or '未指定'}

【RAG / 类案参考材料】
{rag_ctx}

【时间、地名、人物、交通工具 — 国籍/地区缺省规则】
- 若用户**未**单独指出某时间、地名、人物或交通工具发生在或属于**非中国大陆**，则默认该类**未标注要素**均在中国大陆（中国内陆）发生或具有法律连结；文书计划与后续追问须与此前提一致。
- 若用户**已**明确标注非中国大陆的国籍、地区或连结点，则以用户标注为准，并说明适用法域及原因（如属人、属地、合同连结点等）。

请严格使用下列英文方括号标签输出各块（标签单独占一行，便于解析）：

[PLAN]
**第一句（必须是 [PLAN] 正文的第一行、完整陈述句，二选一）：**
- 若用户未对事件中的时间、地名、人物、交通工具等作**非中国大陆**之特别标注，第一句须为：**基于您的表述，该事件的所有要素均受中国内陆法律管辖。**
- 若用户已作非中国大陆标注，第一句须说明：**涉及的法律法域为……（具体法域/地区），原因是……（与用户表述对应的连结点说明）。**
**从第二行起**，用中文分条写出本案**文书与条款拟制总计划**（4–8 条，可用 1. 2. … 或 •）：阶段、重点、拟产出类型等。

[RAG_SLOT_HINTS]
结合参考材料与常识，列出用户**最可能需要补充**的合同/条款细节（槽位式，• 开头 6–12 条）。

[QUESTIONS]
本轮**槽位填充式**具体问题（3–6 个），编号 1. 2. …，简短可答。

[GUIDANCE_NOTE]
2–4 句中文：说明欢迎逐条补充；若用户希望自己完善全部信息、不需要再被追问，请点击对话区下方的**「忽略建议」**。

[HEURISTIC_STATUS]
只写一词：**CONTINUE**（仍需用户补充或确认）或 **COMPLETE**（诉求已足够清晰、可不再追问）。
"""
    raw = call_openai_chat(
        [
            # 【提示词工程】system 再次强调「标签不可省略」，降低解析失败率（与 user 内标签说明双重约束）。
            {"role": "system", "content": "你只按用户要求的标签块输出，不要省略标签行。"},
            {"role": "user", "content": user},
        ],
        api_key,
        temperature=0.25,
        max_tokens=2200,
    )
    plan = _intake_tag_block(raw, "PLAN")
    slots = _intake_tag_block(raw, "RAG_SLOT_HINTS")
    questions = _intake_tag_block(raw, "QUESTIONS")
    note = _intake_tag_block(raw, "GUIDANCE_NOTE")
    status = _intake_tag_block(raw, "HEURISTIC_STATUS")
    complete = _intake_status_is_complete(status)
    return {
        "raw": raw,
        "plan": plan,
        "slots": slots,
        "questions": questions,
        "note": note,
        "heuristic_complete": complete,
        "rag_context": rag_ctx,
        "round": 1,
    }


def continue_intake_guidance_round(
    api_key,
    objective: str,
    jurisdiction: str,
    plan: str,
    accum_supplementary: str,
    user_reply: str,
    rag_context: str,
    round_num: int,
) -> dict:
    # 【Agent 调用工程】Intake 第 2+ 轮：读入用户回复与累计补充，更新问题列表与是否可结束。
    # 【分块】rag_context 截断 6000 字进入本 prompt，避免首轮 RAG 过长撑爆上下文。
    user = f"""你在做多轮要素引导（当前第 {round_num} 轮）。

【用户原始诉求】
{objective}

【法域】
{jurisdiction or '未指定'}

【已定文书拟制计划】
{plan}

【RAG/类案参考（摘要）】
{rag_context[:6000]}

【用户迄今补充汇总】
{accum_supplementary or '（尚无）'}

【用户本轮回复】
{user_reply or '（空）'}

【法域与事实要素】已定计划首句若已说明中国内陆默认管辖或涉外法域，本轮归纳与追问须与之保持一致；未标注非中国大陆要素时仍默认中国内陆连结。

请输出：

[SYNTHESIS]
简要归纳已知信息 + 仍欠缺的槽位（条列）。

[QUESTIONS]
若仍需追问，列出 2–5 个新问题（编号）；若已足够，只写一行：无。

[GUIDANCE_NOTE]
提醒：若用户希望自己完善、不需再追问，请点击对话区下方的**「忽略建议」**。

[HEURISTIC_STATUS]
**CONTINUE** 或 **COMPLETE**（已足够则 COMPLETE）。
"""
    raw = call_openai_chat(
        [
            {"role": "system", "content": "你只按标签块输出。"},
            {"role": "user", "content": user},
        ],
        api_key,
        temperature=0.2,
        max_tokens=1800,
    )
    syn = _intake_tag_block(raw, "SYNTHESIS")
    questions = _intake_tag_block(raw, "QUESTIONS")
    note = _intake_tag_block(raw, "GUIDANCE_NOTE")
    status = _intake_tag_block(raw, "HEURISTIC_STATUS")
    q = (questions or "").strip()
    complete = _intake_status_is_complete(status)
    # 【提示词工程配套】对 QUESTIONS 做「无 / 极短」启发式，减轻模型误标 CONTINUE 导致的无限追问。
    if not complete:
        trivial = ("无", "无。", "无问题", "没有问题", "暂无", "暂无。")
        if q in trivial or (q.startswith("无") and len(q) < 20):
            complete = True
    return {
        "raw": raw,
        "synthesis": syn,
        "questions": questions,
        "note": note,
        "heuristic_complete": complete,
        "round": round_num,
    }


# 将 intake 产物拼进主工作流 objective，使下游 Agent 1～7 在同一段上下文中看到「计划 + 用户补充」
def compose_workflow_objective(base_objective: str, plan: str, supplementary: str) -> str:
    parts = [base_objective.strip()]
    if plan and plan.strip():
        parts.append("\n【文书拟制计划】\n" + plan.strip())
    if supplementary and supplementary.strip():
        parts.append("\n【经要素引导汇总的用户补充】\n" + supplementary.strip())
    return "\n".join(parts).strip()


def format_intake_plan_assistant_message(data: dict) -> str:
    """将小A 首轮计划与引导格式化为对话气泡正文。"""
    plan = (data.get("plan") or "").strip() or "（暂无）"
    slots = (data.get("slots") or "").strip() or "（暂无）"
    qs = (data.get("questions") or "").strip() or "（暂无）"
    note = (data.get("note") or "").strip()
    parts = [
        "下面是根据您与我的对话整理的**文书拟制计划**与需要您补充的要点。",
        "",
        "#### 文书拟制计划",
        plan,
        "",
        "#### RAG / 类案槽位提示",
        slots,
        "",
        "#### 本轮请补充或回答",
        qs,
    ]
    if note:
        parts.extend(["", note])
    parts.append(
        "\n您可直接在下方输入框回复以作补充；若希望自行完善、**不再需要追问**，请点击 **忽略建议**。"
    )
    if data.get("heuristic_complete"):
        parts.append("\n*当前信息已较充分，您也可以直接启动主工作流。*")
    return "\n".join(parts)


def format_intake_followup_assistant_message(out: dict) -> str:
    """多轮要素引导后续回复，写入对话。"""
    syn = (out.get("synthesis") or "").strip() or "（暂无）"
    qs = (out.get("questions") or "").strip() or "（暂无）"
    note = (out.get("note") or "").strip()
    parts = [
        "#### 对您补充的归纳",
        syn,
        "",
        "#### 本轮追问",
        qs,
    ]
    if note:
        parts.extend(["", note])
    parts.append(
        "\n可继续直接回复；不需要再追问时请点 **忽略建议**。"
    )
    if out.get("heuristic_complete"):
        parts.append("\n*要素已较充分，可启动主工作流。*")
    return "\n".join(parts)


# Intake 与主工作流之间的门禁：跳过引导 / 已生成计划且（已忽略或已 ready）
def intake_workflow_gate_ok(ss) -> bool:
    if ss.get("intake_skipped_all"):
        return True
    if not ss.get("intake_plan_generated"):
        return False
    return bool(ss.get("intake_ignored") or ss.get("intake_ready"))


# 清空 intake 相关 session 键；用于「重置引导」按钮，避免陈旧 plan 污染新一轮
def reset_intake_session_state():
    keys = [
        "intake_skipped_all",
        "intake_plan_generated",
        "intake_ignored",
        "intake_ready",
        "intake_phase_round",
        "intake_drafting_plan",
        "intake_rag_hints",
        "intake_guidance_note",
        "intake_questions_md",
        "intake_synthesis_md",
        "intake_supplementary_accum",
        "intake_rag_context",
    ]
    for k in keys:
        st.session_state.pop(k, None)


# 写入历史记录时需快照的 intake 键（与 reset_intake_session_state 一致）
MEMORY_INTAKE_SNAPSHOT_KEYS = [
    "intake_skipped_all",
    "intake_plan_generated",
    "intake_ignored",
    "intake_ready",
    "intake_phase_round",
    "intake_drafting_plan",
    "intake_rag_hints",
    "intake_guidance_note",
    "intake_questions_md",
    "intake_synthesis_md",
    "intake_supplementary_accum",
    "intake_rag_context",
]


# --- 5b) 用户记忆（按「记忆空间」目录隔离；落盘于 user_memory/<slug>/）---
# 知识：Streamlit 无内置账号体系，用侧边栏「记忆空间名称」区分存档；请勿与他人共享该名称
def _safe_memory_slug(label: str) -> str:
    s = re.sub(r"[^\w\u4e00-\u9fff\-.]", "_", (label or "").strip())
    s = re.sub(r"_+", "_", s).strip("._-")[:64]
    return s or "default"


def memory_user_base_dir(slug_label: str) -> Path:
    d = _USER_MEMORY_ROOT / _safe_memory_slug(slug_label)
    d.mkdir(parents=True, exist_ok=True)
    return d


def _memory_index_path(slug_label: str) -> Path:
    return memory_user_base_dir(slug_label) / "sessions_index.json"


def list_memory_session_index(slug_label: str) -> list:
    p = _memory_index_path(slug_label)
    if not p.is_file():
        return []
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, OSError):
        return []


def _write_memory_index(slug_label: str, rows: list) -> None:
    rows = sorted(rows, key=lambda r: r.get("updated_at") or "", reverse=True)
    p = _memory_index_path(slug_label)
    p.write_text(json.dumps(rows[:200], ensure_ascii=False, indent=2), encoding="utf-8")


def _memory_session_dir(slug_label: str, session_id: str) -> Path:
    return memory_user_base_dir(slug_label) / session_id


def load_memory_session(slug_label: str, session_id: str) -> Optional[dict]:
    sd = _memory_session_dir(slug_label, session_id)
    meta_p = sd / "session.json"
    if not meta_p.is_file():
        return None
    try:
        meta = json.loads(meta_p.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None
    corp_p = sd / "corpus.json"
    corpus = []
    if corp_p.is_file():
        try:
            corpus = json.loads(corp_p.read_text(encoding="utf-8"))
            if not isinstance(corpus, list):
                corpus = []
        except (json.JSONDecodeError, OSError):
            corpus = []
    meta["uploaded_texts"] = corpus
    return meta


def save_memory_session(
    slug_label: str,
    session_id: str,
    meta: dict,
    uploaded_texts: list,
) -> None:
    sd = _memory_session_dir(slug_label, session_id)
    sd.mkdir(parents=True, exist_ok=True)
    to_write = {k: v for k, v in meta.items() if k != "uploaded_texts"}
    (sd / "session.json").write_text(
        json.dumps(to_write, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (sd / "corpus.json").write_text(
        json.dumps(uploaded_texts or [], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    idx = list_memory_session_index(slug_label)
    title = (meta.get("title") or meta.get("objective_base") or "未命名")[:120]
    row = {
        "session_id": session_id,
        "title": title,
        "updated_at": meta.get("updated_at") or datetime.now().isoformat(),
    }
    idx = [r for r in idx if r.get("session_id") != session_id]
    idx.insert(0, row)
    _write_memory_index(slug_label, idx)


def append_memory_followup(slug_label: str, session_id: str, user_text: str, assistant_text: str) -> None:
    meta = load_memory_session(slug_label, session_id)
    if not meta:
        return
    now = datetime.now().isoformat()
    fu = meta.get("followup_chat")
    if not isinstance(fu, list):
        fu = []
    fu.append({"role": "user", "content": user_text, "ts": now})
    fu.append({"role": "assistant", "content": assistant_text, "ts": now})
    meta["followup_chat"] = fu
    meta["updated_at"] = now
    ut = meta.get("uploaded_texts") or []
    save_memory_session(slug_label, session_id, meta, ut)


def apply_loaded_memory_to_session_state(ss, data: dict) -> None:
    """将磁盘会话恢复到 Streamlit session_state（含表单与引导状态）。"""
    ss["objective_input"] = data.get("objective_base") or data.get("effective_objective") or ""
    ss["jurisdiction_input"] = data.get("jurisdiction") or ""
    style = data.get("firm_style") or "Balanced (Legal but Readable)"
    styles = ["Plain English", "Legal Formal", "Balanced (Legal but Readable)"]
    ss["style_input"] = style if style in styles else styles[2]
    ss["refinement_slider"] = int(data.get("num_refinements") or 2)
    intake = data.get("intake") or {}
    for k, v in intake.items():
        ss[k] = v
    corp = data.get("uploaded_texts") or []
    ss["memory_active_corpus"] = corp if isinstance(corp, list) else []
    ss["memory_viewing_sid"] = data.get("session_id")
    ss["memory_viewing_title"] = data.get("title") or ""


def run_memory_followup_turn(
    api_key: str, slug_label: str, session_id: str, user_message: str
) -> str:
    """基于已保存目标与定稿条款做后续单轮问答（压缩上下文，不 replay 全 Agent 轨迹）。"""
    # 【Agent 调用工程】历史追问用「单轮大 user」而非多轮 chat history，减少 token 与编排复杂度。
    # 【提示词工程】system 定角色+免责声明；user 打包 goal/clause/最近追问摘要/新问题，属上下文压缩（非全量 replay）。
    # 【分块】goal/clause/每条 followup 在组包时截断，防止超长条款撑爆窗口。
    meta = load_memory_session(slug_label, session_id)
    if not meta:
        return "无法读取该条历史记录。"
    goal = (meta.get("effective_objective") or meta.get("objective_base") or "")[:6000]
    clause = (meta.get("final_clause") or "")[:12000]
    prev = meta.get("followup_chat") or []
    if not isinstance(prev, list):
        prev = []
    hist_lines = []
    for m in prev[-12:]:
        if isinstance(m, dict) and m.get("role") in ("user", "assistant"):
            tag = "用户" if m["role"] == "user" else "助手"
            hist_lines.append(f"{tag}: {(m.get('content') or '')[:4000]}")
    hist_blob = "\n".join(hist_lines) if hist_lines else "（尚无）"
    user_blob = f"""【用户原始起草目标与上下文（节选）】
{goal}

【当时生成的条款稿（节选）】
{clause}

【此前后续追问摘要】
{hist_blob}

【用户本轮问题】
{user_message}
"""
    messages = [
        {
            "role": "system",
            # 【提示词工程】与 Intake 同角色轴线，强调可执行 + 非法律意见，降低胡编风险表述。
            "content": (
                "你是资深合同与法律文书辅助顾问。根据用户给定的起草背景与条款稿回答后续问题；"
                "表述清晰、可执行；声明不构成正式法律意见。"
            ),
        },
        {"role": "user", "content": user_blob},
    ]
    return call_openai_chat(messages, api_key, temperature=0.25, max_tokens=2000)


def _append_workflow_memory(
    memory_events, role: str, content: str, agent: Optional[str] = None
):
    # 【Agent 观测 / 记忆工程】把每步对外展示内容记入列表供落盘；单条上限 80000 字符为存储与 JSON 体积控制。
    if memory_events is None:
        return
    memory_events.append(
        {
            "role": role,
            "agent": agent,
            "content": (content or "")[:80000],
            "ts": datetime.now().isoformat(),
        }
    )


# --- 6a) 文档解析：工具函数，供主工作流「资料与研究 Agent」与 Intake 上传共用 ---
# 【解析粒度】此处仍为「整文件 → 一条 text」；向量切片在 documents_to_vector_chunks（vector_chunking.py）中完成。
# 【overlap】递归分块 RECURSIVE_CHUNK_OVERLAP=180；语义分块块间 SEMANTIC_INTER_CHUNK_OVERLAP=130，见 vector_chunking.py 中「#」注释。
# 输出 [{"filename", "text"}, ...]；RAG 前会转为带 strategy/chunk_id 的 chunk 列表。
def extract_text_from_uploaded_files(uploaded_files):
    """从 Streamlit 上传对象解析 txt/md/docx 为纯文本列表。"""
    texts = []
    failed_files = []
    
    for uploaded in uploaded_files:
        try:
            if uploaded.name.lower().endswith(".docx"):
                # Word
                doc = Document(uploaded)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif uploaded.name.lower().endswith((".txt", ".md")):
                # 
                content = uploaded.getvalue()
                try:
                    text = content.decode("utf-8")
                except:
                    text = content.decode("latin-1")
            else:
                # 
                content = uploaded.getvalue()
                try:
                    text = content.decode("utf-8")
                except:
                    text = content.decode("latin-1")
            
            # 
            if text.strip():
                texts.append({"filename": uploaded.name, "text": text})
            else:
                failed_files.append(f"{uploaded.name} ()")
                
        except Exception as e:
            failed_files.append(f"{uploaded.name} (: {str(e)})")
    
    # 
    if failed_files:
        st.warning(f" \n" + "\n".join([f"- {f}" for f in failed_files]))
    
    return texts


# --- 6b) RAG：检索 Agent（轻量）— 用 LLM 对 **chunk 编号** 做相关性判决，失败则 simple_retrieve 回退 ---
# 【RAG 工程】语料为向量切片后的块列表；排序 prompt 中每条为「块标签 + 前 800 字摘要」，LLM 选最相关块索引（LLM-as-selector）。
# 【分块】输入 texts 每项宜含 filename / text / strategy / chunk_id；返回命中块的完整 text 供下游注入。
# 【Agent 调用】一次 call_openai_chat；异常或解析失败则回退 simple_retrieve（词重叠，仍以 chunk 为单元）。
def ai_enhanced_retrieve(texts, query, api_key, top_k=3):
    """返回若干最相关 chunk {filename, text, ...}；排序靠 LLM 读块摘要，失败则词重叠兜底。"""
    if not texts:
        return []

    # 【RL】仅主工作流幕内且臂为 lexical_only 时跳过 LLM 排序（与 _rl_collect_verdicts 同生命周期，避免 Intake 误用残留臂）。
    if st.session_state.get("_rl_collect_verdicts") and (
        st.session_state.get("_rl_rerank_arm") == "lexical_only"
    ):
        return simple_retrieve(texts, query, top_k)

    doc_snippets = []
    for i, item in enumerate(texts):
        # 【向量切片】每条候选是一个 chunk；仅摘要进入排序 prompt，降低 token，返回时带回整块全文。
        snippet = item["text"][:800]
        _lbl = chunk_label(item)
        doc_snippets.append(f"[{i}] Chunk: {_lbl}\nExcerpt:\n{snippet}")
    
    combined_docs = "\n\n".join(doc_snippets)

    # 【排序提示词】要求模型按语义相关性挑选编号；固定 RELEVANT_INDICES 行便于解析，避免纯子串匹配误伤（如 "1" 命中 "10"）。
    retrieval_prompt = f"""You are a document-relevance ranker for legal / contract drafting support.

## User query (what we need evidence for)
{query}

## Candidate text chunks (vector slices / retrieval units)
Each block starts with an index in square brackets, then a chunk label (file name + chunking strategy + segment id), then an excerpt (~800 characters). The **full chunk text** is available only for indices you select.

{combined_docs}

## Your task
1. Read the query and each excerpt.
2. Choose up to **{top_k}** chunk indices that are **most semantically useful** for the query (clause drafting, risk analysis, or answering questions implied by the query). Prefer direct topical match over incidental keyword overlap.
3. Order your choices from **most relevant to least relevant**.

## Output format (plain text only, no markdown code fences)
Reply with **exactly two lines** in this shape:

RELEVANT_INDICES: <comma-separated integers>
REASON: <one short sentence in English explaining why the top choice fits the query; if NONE, write "No excerpt was sufficiently relevant.">

If **no** document is sufficiently relevant, reply instead with:

RELEVANT_INDICES: NONE
REASON: <brief explanation>

Do not output anything before or after these two lines.
"""
    
    try:
        retrieval_result = call_openai_chat(
            [
                {
                    "role": "system",
                    "content": (
                        "You only output the two-line format requested by the user. "
                        "Indices must be non-negative integers that appear in the candidate list headers."
                    ),
                },
                {"role": "user", "content": retrieval_prompt},
            ],
            api_key,
            temperature=0.1,
            max_tokens=300,
        )

        def _parse_relevant_indices(raw: str, n_docs: int, k: int):
            """Returns (indices, explicit_none). explicit_none True => model said no relevant docs."""
            if not raw:
                return [], False
            t = raw.strip()
            m = re.search(
                r"RELEVANT_INDICES\s*:\s*(.+?)(?:\n|$)",
                t,
                re.IGNORECASE | re.DOTALL,
            )
            if m:
                line = m.group(1).strip().split("\n")[0].strip()
                if re.match(r"^NONE\b", line, re.I):
                    return [], True
                seen = set()
                out = []
                for num in re.findall(r"\d+", line):
                    idx = int(num)
                    if 0 <= idx < n_docs and idx not in seen:
                        seen.add(idx)
                        out.append(idx)
                    if len(out) >= k:
                        break
                if out:
                    return out, False
                return [], False
            # No RELEVANT_INDICES line: bracketed indices in order of first appearance
            seen = set()
            out = []
            for mo in re.finditer(r"\[(\d+)\]", t):
                idx = int(mo.group(1))
                if 0 <= idx < n_docs and idx not in seen:
                    seen.add(idx)
                    out.append(idx)
                if len(out) >= k:
                    break
            return out, False

        relevant_indices, explicit_none = _parse_relevant_indices(
            retrieval_result, len(texts), top_k
        )
        if explicit_none:
            return []
        if relevant_indices:
            return [texts[i] for i in relevant_indices]
        return simple_retrieve(texts, query, top_k)

    except Exception:
        return simple_retrieve(texts, query, top_k)


def simple_retrieve(texts, query, top_k=3):
    """词重叠检索；ai_enhanced_retrieve 失败时调用。"""
    # 【RAG 兜底】无 LLM 排序时：query 与每个 **chunk** 的 text 做词袋交集，排序取 top_k。
    # 【分块】检索单元为向量切片后的块，而非整文件。
    q_words = set([w.lower() for w in query.split() if len(w) > 3])
    scored = []
    
    for item in texts:
        words = set([w.lower().strip('.,;:\n') for w in item['text'].split() if len(w) > 3])
        overlap = len(q_words & words)
        scored.append((overlap, item))
    
    scored.sort(reverse=True, key=lambda x: x[0])
    return [item for score, item in scored[:top_k] if score > 0]


# --- 7) 定稿导出：非生成 Agent，将最终条款与元数据写入 Word ---
def create_docx(clause_text, metadata):
    """
    Create a professional Word document with proper legal formatting
    
    Args:
        clause_text: The clause content
        metadata: Document metadata (timestamp, objective, etc.)
        
    Returns:
        BytesIO: Document binary stream
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title - Professional Legal Document Style
    title = doc.add_heading(level=1)
    title_run = title.add_run("CONTRACT CLAUSE")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Document Information Section
    info_heading = doc.add_heading("DOCUMENT INFORMATION", level=2)
    info_heading.runs[0].font.name = 'Times New Roman'
    info_heading.runs[0].font.size = Pt(12)
    
    # Metadata table-style layout
    info_para = doc.add_paragraph()
    info_para.add_run("Date Generated: ").bold = True
    info_para.add_run(f"{metadata.get('timestamp', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Drafting Objective: ").bold = True
    info_para.add_run(f"{metadata.get('objective', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Jurisdiction: ").bold = True
    info_para.add_run(f"{metadata.get('jurisdiction', 'Not specified')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Drafting Style: ").bold = True
    info_para.add_run(f"{metadata.get('style', 'N/A')}")
    for run in info_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Clause Content Section
    clause_heading = doc.add_heading("CLAUSE PROVISIONS", level=2)
    clause_heading.runs[0].font.name = 'Times New Roman'
    clause_heading.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing before clause
    
    # Add clause text with proper formatting (remove all markdown and LaTeX artifacts)
    import re
    
    # Function to clean LaTeX formatting
    def clean_latex(text):
        # Remove display math delimiters \[ \]
        text = re.sub(r'\\\[(.*?)\\\]', r'\1', text, flags=re.DOTALL)
        # Remove inline math delimiters \( \)
        text = re.sub(r'\\\((.*?)\\\)', r'\1', text, flags=re.DOTALL)
        # Remove $$ $$ delimiters
        text = re.sub(r'\$\$(.*?)\$\$', r'\1', text, flags=re.DOTALL)
        # Remove $ $ delimiters
        text = re.sub(r'\$(.*?)\$', r'\1', text)
        # Replace LaTeX commands with readable text
        text = text.replace(r'\text{', '').replace('}', '')
        text = text.replace(r'\times', '×')
        text = text.replace(r'\%', '%')
        # Replace \frac{a}{b} with (a / b)
        text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1 / \2)', text)
        # Remove remaining backslashes
        text = text.replace('\\', '')
        return text
    
    # Split into lines and process each
    lines = clause_text.split('\n')
    
    for line in lines:
        if not line.strip():
            continue  # Skip empty lines
            
        # Remove markdown formatting
        clean_line = line.replace("**", "").replace("*", "")
        # Remove markdown headers but keep the text
        clean_line = re.sub(r'^#{1,6}\s+', '', clean_line)
        # Clean LaTeX formatting
        clean_line = clean_latex(clean_line)
        
        # Check if this looks like a header (was markdown header or is short and looks like title)
        if line.startswith('#'):
            # This was a markdown header - make it a sub-heading in Word
            heading_para = doc.add_paragraph(clean_line)
            heading_para.runs[0].bold = True
            heading_para.runs[0].font.name = 'Times New Roman'
            heading_para.runs[0].font.size = Pt(11)
            heading_para.paragraph_format.space_before = Pt(6)
        else:
            # Regular paragraph
            clause_para = doc.add_paragraph(clean_line)
            clause_para.paragraph_format.line_spacing = 1.15
            for run in clause_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing after clause
    
    # Professional Disclaimer
    disclaimer_heading = doc.add_heading("DISCLAIMER", level=2)
    disclaimer_heading.runs[0].font.name = 'Times New Roman'
    disclaimer_heading.runs[0].font.size = Pt(12)
    
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "This document has been generated using artificial intelligence technology "
        "and is provided for informational purposes only. It does not constitute "
        "legal advice, and should not be relied upon as such. Users should consult "
        "with qualified legal professionals before using any content from this document "
        "in actual legal agreements or contracts. The creators and distributors of this "
        "tool disclaim all liability for any damages arising from the use of this document."
    )
    disclaimer.paragraph_format.line_spacing = 1.15
    for run in disclaimer.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.italic = True
    
    # Save to memory
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio


# =============================================================================
# 8) 主多 Agent 合同条款流水线（顺序 DAG + 每步「生成-反思-replan」包裹）
# -----------------------------------------------------------------------------
# 【Agent 工程】7 个固定角色 + N 轮 Review：前序输出写入后序 user prompt，形成单向数据 DAG（无并行分支）。
# 【提示词工程】每步英文 task 描述 + system 定角色；user 内嵌 objective / 上传摘要 / 检索摘录 / 约束等，属 in-context 组装。
# 【RAG】资料研究步：有语料则总结 + ai_enhanced_retrieve(objective)；ReAct 内 rag_retrieve 可多次按 query 取片段。
# 【RL】幕始 ε-贪心选臂（LLM 重排 vs 词重叠）；幕末用 Critic 通过率更新 bandit（见 rl_bandit.py）。
# 【向量切片】combined_preview 按 chunk 展示；evidence 每命中块 1000 字；排序 prompt 内每块摘要 800 字；见 vector_chunking.py 策略说明。
# UI：st.chat_message 模拟多角色；step_key 对齐可观测性与记忆落盘。
# =============================================================================
def run_clause_agent_workflow(
    api_key,
    objective,
    jurisdiction,
    uploaded_files,
    firm_style,
    num_refinements,
    preloaded_texts: Optional[list] = None,
    memory_events: Optional[list] = None,
):
    """主工作流入口：7 个固定阶段 + num_refinements 次审查；全程重置 ai_call_count。"""
    st.session_state.ai_call_count = 0

    # 【RL】为本幕选择检索排序臂，并开启 Critic 奖励采集；仅在本函数运行期间 _rl_collect_verdicts 为 True。
    st.session_state["_rl_episode_verdicts"] = []
    st.session_state["_rl_collect_verdicts"] = True
    st.session_state["_rl_rerank_arm"] = get_rerank_bandit(_PROJECT_ROOT).select_arm()
    _rl_arm_label = (
        "LLM 文档重排序"
        if st.session_state.get("_rl_rerank_arm") == "llm_rerank"
        else "词重叠检索（省一次排序模型调用）"
    )

    progress_bar = st.progress(0)
    status_text = st.empty()
    total_steps = 7 + num_refinements
    step_i = 0

    # 【提示词工程 / UX】编排说明面向终端用户，同步写入 memory_events；不调 LLM，不进入 ReAct。
    orch_msg = (
        "已启动合同条款多 Agent 工作流。每步生成阶段采用 **ReAct 工具循环**（思考与可选检索等工具，直至提交最终答案），"
        "随后由**验证与反思 Agent**检查是否正确、完整、可行；未通过则 **meta-prompt replan** 并重新进入 ReAct。"
        "通过步骤会输出**正则清洗后的固定格式文档**块。"
        f"流程：目标解析 → 资料与研究 → 风险约束 → 条款起草 → 审查优化（{num_refinements} 轮）→ 定稿导出 → 质量评估。\n\n"
        f"**本幕 RL（Bandit）检索策略**：{_rl_arm_label}（幕末按 Critic 通过率更新）。"
    )

    # --- 协调 Agent（Orchestrator）：仅说明与进度，不调 LLM；对应多 Agent 系统中的「编排/元层」---
    with st.chat_message("assistant", avatar="🧭"):
        st.markdown("**协调 Agent**")
        st.markdown(orch_msg)
    _append_workflow_memory(
        memory_events,
        "assistant",
        "**协调 Agent**\n\n" + orch_msg,
        "协调 Agent",
    )

    # --- Agent 1 目标解析：将自然语言诉求结构化为 5 类洞察，供下游对齐「成功标准」---
    # 知识：属于需求澄清类 Agent；输出应短而可执行，避免与后文「起草」重复冗长
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 目标解析 Agent 工作中…")
    progress_bar.progress(step_i / total_steps)
    with st.chat_message("assistant", avatar="📋"):
        st.markdown("**目标解析 Agent**")
        with st.spinner("解析起草目标（含验证 / 必要时 replan）…"):
            # 【提示词工程】五维结构化输出 → 下游对齐「成功标准」；英文明示降低与起草步语言混杂。
            # 【Agent】react_ctx texts 空 → RAG 工具仅返回无文档，避免误用上传材料做目标解析。
            base_user_1 = f"""Based on the drafting request below, produce 5 numbered insights:
1) Core commercial/legal intent
2) Key parties/obligations implied
3) Typical risks to watch
4) Jurisdiction hooks (if any)
5) Success criteria for the clause

Drafting objective: {objective}
Jurisdiction: {jurisdiction or 'Not specified'}
Style: {firm_style}
"""
            # 以下各 agent_call 共性：task_summary / reflect_context 专供 Critic；temperature/max_tokens 控制本步生成风格与长度。
            c1, f1, _, rec1, rp1 = agent_call_with_verify_and_replan(
                api_key,
                step_key="OBJECTIVE_ANALYSIS",
                step_label="目标解析 Agent",
                task_summary="Produce 5 numbered insights from drafting objective, jurisdiction, and style.",
                reflect_context=f"objective={objective[:800]}; jurisdiction={jurisdiction}; style={firm_style}",
                system_prompt="You are a senior contracts lawyer. Be concise and structured.",
                build_user_content=lambda meta, b=base_user_1: b + (f"\n\n{meta}" if meta else ""),
                react_ctx={"texts": []},
                max_tokens=1200,
                max_replans=2,
            )
        streamlit_show_agent_verification(
            "目标解析 Agent", c1, f1, rec1, rp1, main_render=lambda t: st.markdown(t)
        )
        _append_workflow_memory(
            memory_events, "assistant", f"**目标解析 Agent**\n\n{c1}", "目标解析 Agent"
        )

    # --- Agent 2 资料与研究：有上传则「摘要 + ai_enhanced_retrieve」；无上传则「法律背景研究」---
    # 知识：RAG 子步骤在此第二次出现（与 intake 的语料池 RAG 独立）；docs_summary 进入约束 Agent
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 资料与研究 Agent 工作中…")
    progress_bar.progress(step_i / total_steps)
    if uploaded_files:
        _raw_docs = extract_text_from_uploaded_files(uploaded_files)
    elif preloaded_texts:
        _raw_docs = [dict(t) for t in preloaded_texts]
    else:
        _raw_docs = []
    # 【向量切片】.md/.txt → 递归字符分块；.docx → 语义分块（段落嵌入合并）；供 RAG / ReAct 以 chunk 为检索单元。
    texts = documents_to_vector_chunks(_raw_docs, api_key) if _raw_docs else []
    docs_summary = ""
    retrieved = []

    with st.chat_message("assistant", avatar="📚"):
        st.markdown("**资料与研究 Agent**")
        if texts:
            # 总结 prompt 只展示部分 chunk 摘要，避免超长；策略标签见 chunk_label。
            _prev_n = min(len(texts), 36)
            combined_preview = "\n\n".join(
                [
                    f"--- {chunk_label(t)} ---\n{t['text'][:1200]}"
                    for t in texts[:_prev_n]
                ]
            )
            if len(texts) > _prev_n:
                combined_preview += f"\n\n… 另有 {len(texts) - _prev_n} 个向量块未展开 …"
            with st.expander("上传文件预览"):
                st.text_area("内容摘要预览", combined_preview, height=200, key="docs_preview_agent")
            with st.spinner("总结上传材料（含验证 / 必要时 replan）…"):
                # 【提示词工程】三列表（模式/定义项/缺口）强制模型做「材料—目标」对齐，便于约束 Agent 引用。
                base_user_docs = f"""Summarize the uploaded materials for this clause objective. List:
1) Reusable language patterns
2) Defined terms or caps
3) Gaps vs the objective

Objective: {objective}

Materials:
{combined_preview}
"""
                cd, fd, rawd, recd, rpd = agent_call_with_verify_and_replan(
                    api_key,
                    step_key="DOCS_SUMMARY",
                    step_label="资料与研究 Agent · 材料总结",
                    task_summary="Summarize uploads: patterns, defined terms/caps, gaps vs objective.",
                    reflect_context=(
                        f"objective={objective[:500]}; chunks={len(texts)}; "
                        f"files={len({t.get('filename','') for t in texts})}"
                    ),
                    system_prompt="You summarize legal/reference documents clearly.",
                    build_user_content=lambda meta, b=base_user_docs: b + (f"\n\n{meta}" if meta else ""),
                    react_ctx={"texts": texts},
                    max_tokens=1400,
                    max_replans=2,
                )
            streamlit_show_agent_verification(
                "资料与研究 Agent · 材料总结",
                cd,
                fd,
                recd,
                rpd,
                main_render=lambda t: st.markdown(t),
            )
            _append_workflow_memory(
                memory_events,
                "assistant",
                f"**资料与研究 Agent · 材料总结**\n\n{cd}",
                "资料与研究 Agent",
            )
            docs_summary = rawd
            with st.spinner("检索相关片段…"):
                # 【RAG】以整条 objective 为 query，与 Intake 一致；top_k=3 控制进入证据块的文档条数。
                retrieved = ai_enhanced_retrieve(texts, objective, api_key, top_k=3)
            if retrieved:
                st.success(f"已检索 {len(retrieved)} 个相关向量块")
                with st.expander("检索到的片段"):
                    for r in retrieved:
                        st.markdown(f"**{chunk_label(r)}**")
                        # UI 仅展示 500 字；证据块注入约束 Agent 时用更长截断。
                        st.code(r["text"][:500], language="text")
            else:
                st.info("未从上传文件中匹配到高相关片段，后续将主要依赖通用研究结论。")
        else:
            with st.spinner("法律背景研究（含验证 / 必要时 replan）…"):
                # 【提示词工程】无上传时的「广义知识」步；system 再强调非辖区具体法律意见，降低幻觉被当作法条的风险。
                base_user_lr = f"""No uploads. Give concise legal background for drafting this clause.
Objective: {objective}
Jurisdiction hint: {jurisdiction or 'unspecified'}

Cover: typical structures, market norms, enforcement pitfalls. Mark assumptions clearly.
"""
                clr, flr, rawlr, reclr, rplr = agent_call_with_verify_and_replan(
                    api_key,
                    step_key="LEGAL_RESEARCH",
                    step_label="资料与研究 Agent · 背景研究",
                    task_summary="Legal background for clause drafting without uploads; note assumptions.",
                    reflect_context=f"objective={objective[:800]}; jurisdiction={jurisdiction}",
                    system_prompt="You provide careful general legal context, not jurisdiction-specific advice.",
                    build_user_content=lambda meta, b=base_user_lr: b + (f"\n\n{meta}" if meta else ""),
                    react_ctx={"texts": []},
                    max_tokens=1400,
                    max_replans=2,
                )
            streamlit_show_agent_verification(
                "资料与研究 Agent · 背景研究",
                clr,
                flr,
                reclr,
                rplr,
                main_render=lambda t: st.markdown(t),
            )
            _append_workflow_memory(
                memory_events,
                "assistant",
                f"**资料与研究 Agent · 背景研究**\n\n{clr}",
                "资料与研究 Agent",
            )
            docs_summary = f"(无上传文件)\n\n{rawlr}"

    # --- Agent 3 风险与约束：在证据块（检索片段）与摘要上生成 Must-haves / Risks / 开放问题---
    # 知识：典型「红队/合规」视角；输出结构化供起草 Agent 直接当约束清单使用
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 风险与约束 Agent 工作中…")
    progress_bar.progress(step_i / total_steps)
    # 证据块按「向量块」注入，标签含分块策略，便于对照检索单元。
    evidence_block = (
        "\n\n".join([f"From {chunk_label(r)}\n{r['text'][:1000]}" for r in retrieved])
        if retrieved
        else "(No retrieved excerpts)"
    )

    with st.chat_message("assistant", avatar="⚖️"):
        st.markdown("**风险与约束 Agent**")
        with st.spinner("分析约束与风险（含验证 / 必要时 replan）…"):
            # 【提示词工程】红队式四段输出（Must-haves/Risks/Enhancements/Open questions）→ 起草 Agent 直接当检查清单。
            base_user_c = f"""Objective: {objective}
Jurisdiction: {jurisdiction or 'Not specified'}
Prior summary/context:
{docs_summary}

Retrieved excerpts:
{evidence_block}

Output sections:
**A. Must-haves** (3–5 bullets)
**B. Risks** (2–3 bullets)
**C. Optional enhancements** (3–5 bullets)
**D. Open questions** for drafter
"""
            cc, fc, rawc, recc, rpc = agent_call_with_verify_and_replan(
                api_key,
                step_key="CONSTRAINTS_RISK",
                step_label="风险与约束 Agent",
                task_summary="Must-haves, risks, optional enhancements, open questions for drafter.",
                reflect_context=f"objective={objective[:600]}; excerpts_len={len(evidence_block)}",
                system_prompt="You identify legal/commercial constraints and risks for contract drafting.",
                build_user_content=lambda meta, b=base_user_c: b + (f"\n\n{meta}" if meta else ""),
                react_ctx={"texts": texts},
                max_tokens=1400,
                max_replans=2,
            )
        streamlit_show_agent_verification(
            "风险与约束 Agent", cc, fc, recc, rpc, main_render=lambda t: st.markdown(t)
        )
        _append_workflow_memory(
            memory_events, "assistant", f"**风险与约束 Agent**\n\n{cc}", "风险与约束 Agent"
        )
        constraints = rawc

    # --- Agent 4 条款起草：在约束下产出可执行条款 + Drafting Notes；regex_clean 拆分正文与说明---
    # 知识：执行型 Agent；current_clause 进入审查循环，需控制 LaTeX/符号（system prompt 已约束）
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 条款起草 Agent 工作中…")
    progress_bar.progress(step_i / total_steps)
    with st.chat_message("assistant", avatar="✍️"):
        st.markdown("**条款起草 Agent**")
        with st.spinner("起草初稿（含验证 / 必要时 replan）…"):
            # 【提示词工程】硬编码「Drafting Notes」标题，便于代码 split 分离条款与说明；禁止 LaTeX 在 system 重申。
            base_user_draft = f"""Draft the clause.

Objective: {objective}
Constraints analysis:
{constraints}
Style: {firm_style}

Requirements:
1) Numbered sub-clauses where helpful
2) Defined terms on first use
3) After the clause, add a section titled exactly:

Drafting Notes
• Note 1: ...
• Note 2: ...
"""
            c_ic, f_ic, raw_ic, rec_ic, rp_ic = agent_call_with_verify_and_replan(
                api_key,
                step_key="DRAFT_INITIAL",
                step_label="条款起草 Agent",
                task_summary="Draft contract clause plus Drafting Notes section; plain text, no LaTeX.",
                reflect_context=f"objective={objective[:500]}; style={firm_style}",
                system_prompt=(
                    "You are an experienced contract lawyer. Use plain text only — no LaTeX. "
                    "For ratios use (A / B)."
                ),
                build_user_content=lambda meta, b=base_user_draft: b + (f"\n\n{meta}" if meta else ""),
                react_ctx={"texts": texts},
                max_tokens=1500,
                max_replans=2,
            )
        with st.expander("条款起草 Agent · 验证与反思", expanded=rp_ic > 0):
            for rec in rec_ic:
                st.markdown(f"**轮次 {rec['attempt']} · 结论：`{rec['verdict']}`**")
                st.markdown(rec["raw_reflection"])
                rt = rec.get("react_trace")
                if rt:
                    with st.expander(
                        f"ReAct 工具轨迹（第 {rec['attempt']} 次尝试）", expanded=False
                    ):
                        for j, ev in enumerate(rt, 1):
                            st.markdown(f"**{j}. `{ev.get('tool')}`**")
                            st.caption(
                                json.dumps(ev.get("args"), ensure_ascii=False)[:1200]
                                or "{}"
                            )
                            st.text((ev.get("observation") or "")[:4000])
            if rp_ic > 0:
                st.info(f"本步骤已 meta-prompt replan {rp_ic} 次。")
        explanation_part = ""
        if "Drafting Notes" in raw_ic:
            parts = raw_ic.split("Drafting Notes", 1)
            clause_part = parts[0].strip()
            explanation_part = parts[1].strip() if len(parts) > 1 else ""
            current_clause = regex_clean_agent_text(clause_part)
        else:
            current_clause = regex_clean_agent_text(c_ic)
        st.markdown("### 初稿条款（供下游审查）")
        st.code(current_clause, language="text")
        if explanation_part:
            with st.expander("起草说明（Drafting Notes）"):
                st.markdown(regex_clean_agent_text(explanation_part))
        with st.expander("条款起草 Agent · 固定格式文档（正则清洗）", expanded=False):
            st.code(f_ic, language="text")
        draft_blob = f"### 初稿条款\n{current_clause}\n\n### Drafting Notes\n{explanation_part or '（无）'}"
        _append_workflow_memory(
            memory_events,
            "assistant",
            f"**条款起草 Agent**\n\n{draft_blob}",
            "条款起草 Agent",
        )

    # --- Agent 5 审查优化（循环 num_refinements 轮）：每轮解析 [Revised Clause] / [Revision Notes]---
    # 知识：Critic-Editor 循环；轮数由用户 slider 控制，平衡质量与 API 成本
    st.markdown("### 审查优化 Agent（多轮）")
    for i in range(num_refinements):
        step_i += 1
        status_text.info(f"进度 {step_i}/{total_steps} · 审查优化 Agent 第 {i + 1} 轮…")
        progress_bar.progress(step_i / total_steps)
        with st.chat_message("assistant", avatar="🔎"):
            st.markdown(f"**审查优化 Agent · 第 {i + 1} 轮**")
            round_n = i + 1
            cc_snapshot = current_clause

            def build_review_user(meta, cc=cc_snapshot, rn=round_n):
                # 【提示词工程】强制 [Revised Clause]/[Revision Notes] 标签，与下游字符串解析耦合；meta 为 replan 追加块。
                base = f"""Review and refine the clause.

**Drafting Objective**: {objective}

**Current Clause**:
{cc}

**Output Format** (follow exactly):

[Revised Clause]
(complete revised text)

[Revision Notes]
- change 1
- change 2

Review round: {rn}
"""
                return base + (f"\n\n{meta}" if meta else "")

            with st.spinner(f"审查与修订 · 第 {round_n} 轮（含验证 / 必要时 replan）…"):
                cr, fr, raw_r, rec_r, rp_r = agent_call_with_verify_and_replan(
                    api_key,
                    step_key=f"REVIEW_ROUND_{round_n}",
                    step_label=f"审查优化 Agent · 第 {round_n} 轮",
                    task_summary="Review clause; output [Revised Clause] and [Revision Notes].",
                    reflect_context=f"objective={objective[:400]}; round={round_n}",
                    system_prompt="You are a professional contract lawyer reviewing and refining clauses.",
                    build_user_content=build_review_user,
                    react_ctx={"texts": texts},
                    max_tokens=1500,
                    max_replans=2,
                )
            with st.expander(
                f"审查优化 Agent · 第 {round_n} 轮 · 验证与反思", expanded=rp_r > 0
            ):
                for rec in rec_r:
                    st.markdown(f"**轮次 {rec['attempt']} · 结论：`{rec['verdict']}`**")
                    st.markdown(rec["raw_reflection"])
                    rt = rec.get("react_trace")
                    if rt:
                        with st.expander(
                            f"ReAct 工具轨迹（第 {rec['attempt']} 次尝试）", expanded=False
                        ):
                            for j, ev in enumerate(rt, 1):
                                st.markdown(f"**{j}. `{ev.get('tool')}`**")
                                st.caption(
                                    json.dumps(ev.get("args"), ensure_ascii=False)[:1200]
                                    or "{}"
                                )
                                st.text((ev.get("observation") or "")[:4000])
                if rp_r > 0:
                    st.info(f"本轮已 meta-prompt replan {rp_r} 次。")
            if "Revised Clause" in raw_r or "[Revised Clause]" in raw_r:
                if "[Revision Notes]" in raw_r:
                    parts = raw_r.split("[Revision Notes]")
                elif "Revision Notes" in raw_r:
                    parts = raw_r.split("Revision Notes")
                else:
                    parts = [raw_r, ""]
                revised_clause = (
                    parts[0]
                    .replace("[Revised Clause]", "")
                    .replace("Revised Clause", "")
                    .strip()
                )
                revised_clause = revised_clause.replace(
                    "(Complete revised clause text here)", ""
                ).strip()
                changes = parts[1].strip() if len(parts) > 1 else ""
                current_clause = regex_clean_agent_text(revised_clause)
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("修订后条款（清洗后）")
                    st.code(current_clause, language="text")
                with c2:
                    st.markdown("修订说明")
                    st.markdown(changes or "_无分项说明_")
            else:
                current_clause = regex_clean_agent_text(cr)
                st.markdown("修订输出（未解析到标记，已整体清洗）")
                st.code(current_clause, language="text")
            with st.expander(
                f"审查优化 · 第 {round_n} 轮 · 固定格式文档（正则清洗）", expanded=False
            ):
                st.code(fr, language="text")
            _append_workflow_memory(
                memory_events,
                "assistant",
                f"**审查优化 Agent · 第 {round_n} 轮**\n\n{current_clause}",
                f"审查优化 Agent · 第 {round_n} 轮",
            )

    # --- Agent 6 定稿：展示最终条款 + create_docx 下载；metadata 写入 objective 等便于归档---
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 定稿与导出…")
    progress_bar.progress(step_i / total_steps)
    with st.chat_message("assistant", avatar="📄"):
        st.markdown("**定稿 Agent**")
        st.code(current_clause, language="text")
        metadata = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "objective": objective,
            "jurisdiction": jurisdiction or "Not specified",
            "style": firm_style,
            "ai_calls": st.session_state.ai_call_count,
        }
        docx_bio = create_docx(current_clause, metadata)
        st.download_button(
            label="下载 Word 文档",
            data=docx_bio,
            file_name=f"AI_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="dl_docx_agent",
        )
        _append_workflow_memory(
            memory_events,
            "assistant",
            f"**定稿 Agent**\n\n{current_clause}",
            "定稿 Agent",
        )

    # --- Agent 7 质量评估：10 维打分 + 优劣势；不回流修改条款（纯评估类 Agent）---
    step_i += 1
    status_text.info(f"进度 {step_i}/{total_steps} · 质量评估 Agent…")
    progress_bar.progress(step_i / total_steps)
    with st.chat_message("assistant", avatar="📊"):
        st.markdown("**质量评估 Agent**")
        cc_final = current_clause
        with st.spinner("评估条款质量（含验证 / 必要时 replan）…"):
            # 【提示词工程】10 维打分 + 总分 + 优劣势，结构化便于人读；本步不回流修改条款（纯评估）。
            base_user_ev = f"""Assess this clause vs the objective.

**Objective**: {objective}
**Final Clause**:
{cc_final}

Use this format:

[Scoring]
1. Objective Achievement: X/10
2. Legal Validity: X/10
3. Language Clarity: X/10
4. Logical Rigor: X/10
5. Enforceability: X/10
6. Risk Control: X/10
7. Professionalism: X/10
8. Completeness: X/10
9. Applicability: X/10
10. Overall Quality: X/10
Total Score: XX/100

[Strengths]
• ...

[Areas for Improvement]
• ...
"""
            ce, fe, _, rece, rpe = agent_call_with_verify_and_replan(
                api_key,
                step_key="QUALITY_EVAL",
                step_label="质量评估 Agent",
                task_summary="Score clause on 10 dimensions; strengths; improvements.",
                reflect_context=f"objective={objective[:500]}; clause_len={len(cc_final)}",
                system_prompt="You are a senior legal expert scoring clause quality.",
                build_user_content=lambda meta, b=base_user_ev: b + (f"\n\n{meta}" if meta else ""),
                react_ctx={"texts": texts},
                max_tokens=1600,
                max_replans=2,
            )
        streamlit_show_agent_verification(
            "质量评估 Agent", ce, fe, rece, rpe, main_render=lambda t: st.markdown(t)
        )
        _append_workflow_memory(
            memory_events, "assistant", f"**质量评估 Agent**\n\n{ce}", "质量评估 Agent"
        )

    status_text.success(f"工作流完成 · 本会话累计 API 调用：{st.session_state.ai_call_count} 次")
    progress_bar.progress(1.0)

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.metric("API 调用次数", st.session_state.ai_call_count)
    with m2:
        st.metric("固定阶段数", 7)
    with m3:
        st.metric("审查轮数", num_refinements)
    with m4:
        _nf = len({t.get("filename", "") for t in texts}) if texts else 0
        st.metric("语料文件数", _nf)
        st.caption(f"向量块数 {len(texts) if texts else 0}")
    st.balloons()

    # 【RL】幕终：用本幕全部 Critic 判定的平均通过率作为 reward，更新所选臂；关闭采集以免 Intake 等误用 lexical 臂。
    st.session_state["_rl_collect_verdicts"] = False
    _rl_arm = st.session_state.pop("_rl_rerank_arm", None)
    _rl_scores = st.session_state.pop("_rl_episode_verdicts", None)
    if _rl_arm and _rl_scores:
        _r = sum(_rl_scores) / len(_rl_scores)
        get_rerank_bandit(_PROJECT_ROOT).update(_rl_arm, _r)

    return {
        "final_clause": current_clause,
        "quality_markdown": ce,
        "corpus_texts": texts,
    }


def execute_main_clause_workflow(api_key: str, uploaded_files) -> bool:
    """
    对话流中自动执行主工作流（原「启动 Agent 合同条款工作流」）。
    成功落盘后重置 intake，下一轮用户发言将重新生成计划。
    """
    ss = st.session_state
    objective = (ss.get("objective_input") or "").strip()
    jurisdiction = (ss.get("jurisdiction_input") or "").strip()
    firm_style = ss.get("style_input") or "Balanced (Legal but Readable)"
    num_refinements = int(ss.get("refinement_slider") or 2)

    if not api_key:
        st.error(
            "无法调用模型：请完成访问验证以使用服务端 API Key，或在侧栏填写 OpenAI API Key。"
        )
        return False
    if not objective or len(objective) < 10:
        st.error("起草目标不足 10 个字符，请先在对话中说明法律诉求。")
        return False
    if not intake_workflow_gate_ok(ss):
        st.error("当前尚不满足启动条件（请完成小A 的要素引导或点「忽略建议」）。")
        return False

    ss["_rl_collect_verdicts"] = False
    _base_for_workflow = objective.strip()
    if ss.get("intake_skipped_all"):
        _base_for_workflow = (
            _base_for_workflow + "\n\n" + WORKFLOW_FACT_ELEMENTS_JURISDICTION_NOTE
        )
    eff_objective = compose_workflow_objective(
        _base_for_workflow,
        "" if ss.get("intake_skipped_all") else ss.get("intake_drafting_plan") or "",
        "" if ss.get("intake_skipped_all") else ss.get("intake_supplementary_accum") or "",
    )

    memory_events: list = []
    user_task_blob = (
        f"【本轮工作流任务】\n{eff_objective}\n\n"
        f"司法辖区：{(jurisdiction or '').strip() or '未填'}\n"
        f"行文风格：{firm_style}\n审查轮数：{num_refinements}"
    )
    _append_workflow_memory(memory_events, "user", user_task_blob, None)

    _pre_corpus = None
    if not uploaded_files:
        _pre_corpus = ss.get("memory_active_corpus")

    workflow_out = run_clause_agent_workflow(
        api_key=api_key,
        objective=eff_objective,
        jurisdiction=(jurisdiction or "").strip(),
        uploaded_files=uploaded_files,
        firm_style=firm_style,
        num_refinements=num_refinements,
        preloaded_texts=_pre_corpus,
        memory_events=memory_events,
    )

    if not workflow_out:
        return False

    _save_slug = _safe_memory_slug((ss.get("memory_user_slug") or "").strip() or "default")
    _new_sid = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
    _now_iso = datetime.now().isoformat()
    _intake_snap = {k: ss.get(k) for k in MEMORY_INTAKE_SNAPSHOT_KEYS}
    _meta = {
        "session_id": _new_sid,
        "title": objective.strip()[:100],
        "objective_base": objective.strip(),
        "effective_objective": eff_objective,
        "jurisdiction": (jurisdiction or "").strip(),
        "firm_style": firm_style,
        "num_refinements": int(num_refinements),
        "intake": _intake_snap,
        "workflow_transcript": memory_events,
        "final_clause": workflow_out.get("final_clause") or "",
        "quality_markdown": workflow_out.get("quality_markdown") or "",
        "followup_chat": [],
        "created_at": _now_iso,
        "updated_at": _now_iso,
    }
    save_memory_session(
        _save_slug,
        _new_sid,
        _meta,
        workflow_out.get("corpus_texts") or [],
    )
    ss.memory_active_corpus = workflow_out.get("corpus_texts") or []
    ss.memory_viewing_sid = _new_sid
    ss.memory_viewing_title = _meta["title"]
    reset_intake_session_state()
    ss.hero_chat_messages.append(
        {
            "role": "assistant",
            "content": "本轮合同条款工作流已完成，结果已保存。您可继续提出新的法律诉求，我将再次为您生成计划并起草。",
        }
    )
    st.success(
        "本次工作流已保存到当前记忆空间，可在主区「记忆空间与历史记录」中再次载入。"
    )
    st.toast("工作流已完成。")
    return True


def render_main_workflow_strip(api_key: Optional[str]) -> None:
    """主区：文书引导状态说明；条款工作流由对话发送自动触发（上传在侧栏）。"""
    st.markdown("---")
    ss = st.session_state
    st.caption(
        "在对话中说明法律诉求与背景后，将**自动生成**文书计划并引导补充；条件满足时**自动启动**合同条款多 Agent 工作流，无需再点按钮。"
        "参考文件请在左侧栏上传。"
    )

    if (
        ss.get("intake_plan_generated")
        and not ss.get("intake_skipped_all")
        and not ss.get("intake_ignored")
    ):
        st.info(
            (ss.get("intake_guidance_note") or "")
            + " **需要自行完善、不再追问时，可点「忽略建议」**（将直接按当前信息启动工作流）。"
        )
        if st.button("忽略建议", type="primary", key="intake_ignore_btn"):
            ss["intake_ignored"] = True
            if api_key:
                execute_main_clause_workflow(api_key, ss.get("upload_input"))
            else:
                ss.hero_chat_messages.append(
                    {
                        "role": "assistant",
                        "content": "需要可用的 API Key 后才能启动工作流。",
                    }
                )
            st.rerun()
    elif ss.get("intake_plan_generated") and not ss.get("intake_skipped_all") and ss.get(
        "intake_ignored"
    ):
        st.caption("已忽略要素追问，工作流将按当前信息运行（若尚未运行请再发送一条消息）。")

    if intake_workflow_gate_ok(ss):
        st.success("**已满足运行条件**：发送下一条对话消息将自动启动条款工作流（若尚未自动启动）。")
    elif not ss.get("intake_skipped_all"):
        st.info(
            "小A 正在通过对话引导您补充要素；**请直接回复对话**，就绪后将自动进入条款起草工作流。"
        )

    if ss.get("ai_call_count") is not None:
        st.caption(f"本会话已累计 API 调用：{ss.ai_call_count} 次")


def render_memory_panel_main(slug_safe: str) -> None:
    """记忆空间与历史载入（原侧栏功能，现置于主区折叠区）。"""
    with st.expander("记忆空间与历史记录", expanded=False):
        st.caption(
            "不同「记忆空间名称」对应服务器上独立文件夹；请勿与他人共享名称以免混档。"
        )
        st.text_input(
            "记忆空间名称",
            key="memory_user_slug",
            help="字母、数字、中文、点或横线；用于区分您的存档目录。",
        )
        if st.button("清除当前载入的语料与查看状态", key="memory_clear_loaded_btn"):
            st.session_state.pop("memory_active_corpus", None)
            st.session_state.pop("memory_viewing_sid", None)
            st.session_state.pop("memory_viewing_title", None)
            st.rerun()
        st.markdown("---")
        st.markdown("**历史记录**（点击载入可回填表单并回顾）")
        _idx_rows = list_memory_session_index(slug_safe)
        if not _idx_rows:
            st.caption("暂无保存记录；完成一次主工作流后会自动存档。")
        else:
            for _row in _idx_rows[:30]:
                _sid = _row.get("session_id") or ""
                _tl = (_row.get("title") or _sid)[:56]
                _when = (_row.get("updated_at") or "")[:19]
                c_a, c_b = st.columns([3, 1])
                with c_a:
                    st.caption(f"{_when}\n{_tl}")
                with c_b:
                    if st.button("载入", key=f"memload_{_sid}"):
                        st.session_state["_memory_do_restore_sid"] = _sid
                        st.rerun()


# =============================================================================
# 9) Streamlit 脚本入口（页面生命周期自上而下执行一次 = 一轮 rerun）
# -----------------------------------------------------------------------------
# 顺序：set_page_config → 全局 CSS → 密码门 st.stop → 侧栏 API/参考上传 → 主区记忆折叠 → Hero/历史对话 → 主区工作流条 → run_clause_agent_workflow
# 知识：session_state 跨 rerun 持久化；按钮点击触发整脚本重跑，故用 pop/flag 传递「仅显示一次」的 UI 状态
# =============================================================================

# Page configuration
st.set_page_config(
    page_title="AI Contract Clause Builder",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 密码门：仅当配置了 BACKEND_PASSWORD 时启用；通过后 backend_verified 解锁后续与托管 OPENAI_API_KEY ---
_backend_password = resolve_backend_password()
if _backend_password and not st.session_state.get("backend_verified"):
    render_access_login_page("AI Contract Clause Builder", _backend_password)
    st.stop()

HERO_ASSISTANT_GREETING = (
    "您好，欢迎使用独属于您的法律文书助手。我是小A，"
    "请告诉我您的法律诉求及相关事件、人物背景，让我为您起草相关的法律文件作为参考哦。"
)

HERO_DIALOG_TITLE = "与小A对话"

# 侧栏顶部展示名 = 新会话默认「记忆空间名称」（同一标识，改此处即可同步）
SIDEBAR_DISPLAY_USER_NAME = "168"

HERO_JURISDICTION_REMINDER = (
    "请您在描述事件过程时，对非中国大陆的时间、地名、人物、交通工具都进行国籍或地区的标明，使回答更准确哦。"
)

# 用户跳过文书计划时写入工作流 objective，与 intake [PLAN] 首句规则对齐
WORKFLOW_FACT_ELEMENTS_JURISDICTION_NOTE = (
    "【事实要素与法域缺省】若用户未单独标注某时间、地名、人物或交通工具属于或发生于非中国大陆，"
    "则默认该等未标注要素发生于中国大陆（中国内陆）并具有中国内陆法律连结，后续分析须与此一致；"
    "若用户已明确标注其他国家/地区或法域，以用户标注为准，并说明所涉法域及连结原因。"
)

# --- 记忆空间 / Hero 对话初始状态 ---
if "memory_user_slug" not in st.session_state:
    st.session_state.memory_user_slug = SIDEBAR_DISPLAY_USER_NAME
if "hero_chat_messages" not in st.session_state:
    st.session_state.hero_chat_messages = [
        {"role": "assistant", "content": HERO_ASSISTANT_GREETING}
    ]
# 起草参数控件已隐藏：用法域/风格/轮数默认值（可在会话状态中改，无前端入口）
if "jurisdiction_input" not in st.session_state:
    st.session_state.jurisdiction_input = "中国大陆"
if "style_input" not in st.session_state:
    st.session_state.style_input = "Balanced (Legal but Readable)"
if "refinement_slider" not in st.session_state:
    st.session_state.refinement_slider = 2

_slug_safe = _safe_memory_slug(
    (st.session_state.get("memory_user_slug") or "").strip() or "default"
)

_restore_sid = st.session_state.pop("_memory_do_restore_sid", None)
if _restore_sid:
    _loaded = load_memory_session(_slug_safe, _restore_sid)
    if _loaded:
        apply_loaded_memory_to_session_state(st.session_state, _loaded)
        _obj_r = (st.session_state.get("objective_input") or "").strip()
        st.session_state.hero_chat_messages = [
            {"role": "assistant", "content": HERO_ASSISTANT_GREETING},
        ]
        if _obj_r:
            st.session_state.hero_chat_messages.append(
                {"role": "user", "content": _obj_r}
            )
        st.toast(
            "已载入历史：表单与引导状态已恢复，上传区可为空（使用记忆中的解析文本）。"
        )
    else:
        st.warning("未找到该条历史，可能已被删除。")

_mem_slug_live = _safe_memory_slug(
    (st.session_state.get("memory_user_slug") or "").strip() or "default"
)
_view_sid = st.session_state.get("memory_viewing_sid")
_vd = None
if _view_sid:
    _vd = load_memory_session(_mem_slug_live, _view_sid)
    if not _vd:
        st.warning("当前查看的历史记录文件不存在或已损坏，已清除查看状态。")
        st.session_state.pop("memory_viewing_sid", None)
        st.session_state.pop("memory_viewing_title", None)
        _view_sid = None

# --- 侧栏：用户名 + API（需在主区「历史追问」前解析 api_key）---
with st.sidebar:
    st.markdown(f"### {SIDEBAR_DISPLAY_USER_NAME}")
    st.markdown("### 法律文书助手")
    st.caption("模型连接与 API 配置")
    api_key = get_api_key_main()
    if _backend_password and st.session_state.get("backend_verified"):
        if st.button(
            "退出验证",
            key="backend_logout_btn",
            help="清除本会话验证状态，需重新输入密码",
        ):
            st.session_state.backend_verified = False
            st.rerun()

    st.markdown("---")
    st.markdown("##### 参考文件（可选）")
    st.caption("支持 .txt / .docx / .md；启动工作流时将一并解析。")
    _uf = st.file_uploader(
        "选择参考文件",
        accept_multiple_files=True,
        type=["txt", "docx", "md"],
        key="upload_input",
        label_visibility="collapsed",
    )
    if _uf:
        st.success(f"已选择 {len(_uf)} 个文件")
    elif st.session_state.get("memory_active_corpus"):
        _nc = len(st.session_state.memory_active_corpus)
        st.info(f"未选新文件时将使用记忆中已载入的 **{_nc}** 个文件解析结果。")

render_memory_panel_main(_mem_slug_live)

# --- 主区：仅保留与「小A」的对话；查看历史时同区展示摘录与继续追问 ---
if _view_sid and _vd:
    st.subheader("历史记录回顾与继续提问")
    st.info(HERO_JURISDICTION_REMINDER)
    st.caption(
        f"正在查看：**{_vd.get('title') or _view_sid}** · "
        "若未重新上传文件，主工作流将优先使用记忆中保存的参考文本。"
    )
    with st.expander("工作流对话摘录（各 Agent 输出）", expanded=True):
        for ev in _vd.get("workflow_transcript") or []:
            if not isinstance(ev, dict):
                continue
            role = ev.get("role") or "assistant"
            content = (ev.get("content") or "")[:50000]
            agent = ev.get("agent") or ""
            av = "🤖" if role == "assistant" else "👤"
            with st.chat_message(role, avatar=av):
                if agent:
                    st.caption(agent)
                st.markdown(content)
    _fu = _vd.get("followup_chat") or []
    if _fu:
        with st.expander("针对该记录的后续追问", expanded=True):
            for m in _fu:
                if not isinstance(m, dict):
                    continue
                rr = m.get("role") or ""
                cc = (m.get("content") or "")[:12000]
                lbl = "您" if rr == "user" else "模型"
                st.markdown(f"**{lbl}**：{cc}")
    if api_key:
        _cq = st.chat_input(
            "针对该条历史向模型继续提问…",
            key="memory_continue_chat_input",
        )
        if _cq and _cq.strip():
            with st.spinner("模型思考中…"):
                _ans = run_memory_followup_turn(
                    api_key, _mem_slug_live, _view_sid, _cq.strip()
                )
            append_memory_followup(_mem_slug_live, _view_sid, _cq.strip(), _ans)
            st.toast("已保存本轮追问与回复。")
            st.rerun()
    else:
        st.caption("配置 API Key 或通过访问验证后，可使用下方输入框继续追问。")
else:
    st.subheader(HERO_DIALOG_TITLE)
    st.info(HERO_JURISDICTION_REMINDER)
    for msg in st.session_state.hero_chat_messages:
        _av = "🤖" if msg["role"] == "assistant" else "👤"
        with st.chat_message(msg["role"], avatar=_av):
            st.markdown(msg["content"])
    render_main_workflow_strip(api_key)
    _hin = st.chat_input(
        "向小A描述您的法律诉求、事件与人物背景…",
        key="hero_chat_main",
    )
    if _hin and _hin.strip():
        _t = _hin.strip()
        st.session_state.hero_chat_messages.append({"role": "user", "content": _t})
        st.session_state.objective_input = "\n\n".join(
            m["content"]
            for m in st.session_state.hero_chat_messages
            if m["role"] == "user"
        )
        _ss = st.session_state
        _ufs = _ss.get("upload_input")

        if not api_key:
            st.error("需要可用的 API Key：完成访问验证或在侧栏填写后，再发送消息以启动流程。")
            st.rerun()

        _obj = (_ss.get("objective_input") or "").strip()
        if len(_obj) < 10:
            st.warning("请至少输入约 **10 个字**，说明法律诉求、事件或人物背景。")
            st.rerun()

        if not (_ss.get("jurisdiction_input") or "").strip():
            _ss["jurisdiction_input"] = "中国大陆"

        # 历史存档里若仅有「跳过引导」而无计划快照，新发言时改走自动生成计划，避免门禁卡死或误连点工作流
        if _ss.get("intake_skipped_all") and not _ss.get("intake_plan_generated"):
            _ss.pop("intake_skipped_all", None)

        # 首轮：尚无文书计划 → 自动生成计划与引导
        if not _ss.get("intake_plan_generated") and not _ss.get("intake_skipped_all"):
            _ss["_rl_collect_verdicts"] = False
            with st.spinner("小A 正在根据您的说明生成文书拟制计划并结合材料做 RAG 联想…"):
                _mc = _ss.get("memory_active_corpus") or []
                _pre_intake = _mc if _mc and not _ufs else None
                data = generate_intake_plan_and_guidance(
                    api_key,
                    _obj,
                    (_ss.get("jurisdiction_input") or "").strip(),
                    _ufs,
                    preloaded_texts=_pre_intake,
                )
            _ss["intake_skipped_all"] = False
            _ss["intake_plan_generated"] = True
            _ss["intake_ignored"] = False
            _ss["intake_ready"] = bool(data["heuristic_complete"])
            _ss["intake_phase_round"] = 1
            _ss["intake_drafting_plan"] = data["plan"]
            _ss["intake_rag_hints"] = data["slots"]
            _ss["intake_questions_md"] = data["questions"]
            _ss["intake_guidance_note"] = data["note"]
            _ss["intake_synthesis_md"] = ""
            _ss["intake_rag_context"] = data["rag_context"]
            _ss["intake_supplementary_accum"] = ""
            _ss.hero_chat_messages.append(
                {
                    "role": "assistant",
                    "content": format_intake_plan_assistant_message(data),
                }
            )
            if _ss.get("intake_ready"):
                execute_main_clause_workflow(api_key, _ufs)
            st.rerun()

        elif (
            api_key
            and _ss.get("intake_plan_generated")
            and not _ss.get("intake_skipped_all")
            and not _ss.get("intake_ignored")
            and not _ss.get("intake_ready")
        ):
            pr = int(_ss.get("intake_phase_round") or 1)
            accum = _ss.get("intake_supplementary_accum") or ""
            chunk = f"\n--- 第{pr}轮用户补充 ---\n{_t}\n"
            new_accum = accum + chunk
            if pr >= INTAKE_MAX_GUIDE_ROUNDS:
                _ss["intake_supplementary_accum"] = new_accum
                _ss["intake_ready"] = True
                _ss.hero_chat_messages.append(
                    {
                        "role": "assistant",
                        "content": "已达到引导轮数上限，您在本轮的补充已保存，将自动进入条款工作流。",
                    }
                )
                execute_main_clause_workflow(api_key, _ufs)
            else:
                with st.spinner("小A 正在根据您的补充整理归纳与追问…"):
                    out = continue_intake_guidance_round(
                        api_key,
                        (_ss.get("objective_input") or "").strip(),
                        (_ss.get("jurisdiction_input") or "").strip(),
                        _ss.get("intake_drafting_plan") or "",
                        new_accum,
                        _t,
                        _ss.get("intake_rag_context") or "",
                        pr + 1,
                    )
                _ss["intake_supplementary_accum"] = new_accum
                _ss["intake_synthesis_md"] = out["synthesis"]
                _ss["intake_questions_md"] = out["questions"]
                _ss["intake_guidance_note"] = out["note"]
                _ss["intake_ready"] = bool(out["heuristic_complete"])
                _ss["intake_phase_round"] = out["round"]
                _ss.hero_chat_messages.append(
                    {
                        "role": "assistant",
                        "content": format_intake_followup_assistant_message(out),
                    }
                )
                if _ss.get("intake_ready"):
                    execute_main_clause_workflow(api_key, _ufs)
            st.rerun()

        elif api_key and intake_workflow_gate_ok(_ss):
            execute_main_clause_workflow(api_key, _ufs)
            st.rerun()
        else:
            st.rerun()
